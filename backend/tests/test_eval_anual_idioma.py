"""El informe final se redacta en el idioma del CA, y ese idioma viaja DENTRO del borrador.

Importa que viaje en el borrador y no solo en la sesión: el Word que descarga el advisee se
regenera desde lo guardado en Notion (`word_desde_borrador`), donde ya no hay sesión local.
Si el idioma no viajara ahí, el CA portugués haría el informe en portugués y su advisee se
bajaría un Word en español.
"""

import pytest
from docx import Document

from backend import eval_anual_sesion as ea


@pytest.fixture(autouse=True)
def entorno_aislado(tmp_path, monkeypatch):
    monkeypatch.setattr(ea.config, "CARPETA_WEB", str(tmp_path))


def _borrador(idioma):
    sesion = {
        "advisee": "Begoña Garcia",
        "ca": "Ana Hernanz",
        "cargo": "Associate",
        "idioma": idioma,
        "emp_data": {"empleado": "Begoña Garcia", "objetivos": []},
        "areas": {},
    }
    return ea._construir_borrador(sesion)


def _texto_docx(ruta):
    doc = Document(ruta)
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)


def test_el_borrador_lleva_dentro_el_idioma_del_ca():
    assert _borrador("pt")["idioma"] == "pt"
    assert _borrador("en")["idioma"] == "en"


def test_idioma_desconocido_cae_a_espanol():
    """Nunca debe colarse un idioma inventado: normalizar_idioma lo pliega a 'es'."""
    assert _borrador("klingon")["idioma"] == "es"


# En español la etiqueta es la de la dimensión tal cual (la de Notion, o la fija si
# Notion no responde); en los demás idiomas se traduce por slug. Ya no existe la
# variante corta "Gestión proyecto": el título del criterio en Notion es el único.
@pytest.mark.parametrize("idioma,fecha_mes,dimension", [
    ("es", "Julio", "Gestión del proyecto"),
    ("en", "July", "Project management"),
    ("pt", "Julho", "Gestão de projeto"),
])
def test_fecha_y_dimensiones_del_borrador_en_el_idioma_del_ca(idioma, fecha_mes, dimension):
    borr = _borrador(idioma)
    assert borr["fecha"].startswith(fecha_mes), borr["fecha"]
    assert borr["dimensiones"][0]["etiqueta"] == dimension


def test_la_clave_de_dimension_no_se_traduce_nunca():
    """`clave` es lo que identifica la fila en Notion y en los merges: traducirla rompería
    los informes ya guardados. Solo `etiqueta` (lo que se pinta) cambia de idioma."""
    for idioma in ("es", "en", "pt"):
        claves = [d["clave"] for d in _borrador(idioma)["dimensiones"]]
        assert claves[0] == "gestion_proyecto"
        assert all(c == c.lower() and " " not in c for c in claves)


@pytest.mark.parametrize("idioma,esperados,intrusos", [
    ("es", ["EVALUACIÓN ANUAL", "Empleado", "Posición actual", "PROMOCIÓN"],
           ["ANNUAL EVALUATION", "AVALIAÇÃO ANUAL"]),
    ("en", ["ANNUAL EVALUATION", "Employee", "Current position", "PROMOTION"],
           ["EVALUACIÓN ANUAL", "AVALIAÇÃO ANUAL"]),
    ("pt", ["AVALIAÇÃO ANUAL", "Colaborador", "Posição atual", "PROMOÇÃO"],
           ["EVALUACIÓN ANUAL", "ANNUAL EVALUATION"]),
])
def test_word_del_advisee_sale_en_el_idioma_en_que_lo_redacto_el_ca(idioma, esperados, intrusos):
    """Camino real del advisee: regenerar el Word desde el borrador de Notion, sin sesión."""
    ruta = ea.word_desde_borrador(_borrador(idioma), f"informe_final_test_{idioma}.docx")
    txt = _texto_docx(ruta)
    for e in esperados:
        assert e in txt, f"falta '{e}' en el Word {idioma}"
    for i in intrusos:
        assert i not in txt, f"se ha colado '{i}' en el Word {idioma}"


def test_informe_antiguo_sin_idioma_sigue_saliendo_en_espanol():
    """Los borradores guardados antes de esto no traen `idioma`: deben seguir saliendo en
    español, que es como se generaron, y no reventar."""
    borr = _borrador("pt")
    del borr["idioma"]
    ruta = ea.word_desde_borrador(borr, "informe_final_test_legacy.docx")
    assert "EVALUACIÓN ANUAL" in _texto_docx(ruta)