"""El anexo de Fuentes/Evidencia del informe final es SOLO para el CA.

El anexo lista, cita a cita, el texto literal de cada evaluación y el nombre de quien la
escribió. Enseñárselo al evaluado rompe el anonimato de los evaluadores, que es lo que
sostiene que la gente escriba con franqueza. Estos tests fijan la frontera por los dos lados:
el documento que se le genera al advisee no lleva anexo, y el que sí lo lleva (la copia con
timestamp que publica el CA) no se le sirve aunque pida el nombre exacto del fichero.
"""

import pytest
from docx import Document

from backend import eval_anual_sesion as ea
from backend import skill_informes_anual as sk
from backend.api import files as files_router

FUENTES = {
    "E1": {
        "tipo": "evaluacion",
        "label": "Eval anual 2025",
        "evaluador": "Marta Ruiz",
        "texto": "Le cuesta delegar y se satura en los picos.",
    },
}


def _texto_docx(ruta):
    doc = Document(ruta)
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.extend(celda.text for celda in fila.cells)
    return "\n".join(partes)


@pytest.fixture(autouse=True)
def carpeta_aislada(tmp_path, monkeypatch):
    monkeypatch.setattr(ea.config, "CARPETA_WEB", str(tmp_path))
    monkeypatch.setattr(sk.config, "CARPETA_WEB", str(tmp_path))


def _borrador():
    sesion = {
        "advisee": "Juan Perez",
        "ca": "Carlos CA",
        "cargo": "Associate",
        "idioma": "es",
        "emp_data": {"empleado": "Juan Perez", "objetivos": []},
        "areas": {},
    }
    return ea._construir_borrador(sesion)


# ── El documento del advisee no lleva anexo ──────────────────────────────────

def test_el_word_del_advisee_no_lleva_anexo_de_fuentes():
    """Camino real del advisee: regenerar desde el borrador de Notion."""
    ruta = ea.word_desde_borrador(_borrador(), "informe_final_Juan_Perez.docx")
    txt = _texto_docx(ruta)
    assert "FUENTES" not in txt.upper()
    assert "Marta Ruiz" not in txt


def test_word_desde_borrador_no_filtra_fuentes_aunque_notion_las_traiga():
    """El dict vacío no basta como defensa: si algún día el borrador guardado trajera
    `_fuentes`, el advisee no debe verlas igualmente. Lo garantiza incluir_fuentes=False."""
    borr = _borrador()
    borr["_fuentes"] = FUENTES
    ruta = ea.word_desde_borrador(borr, "informe_final_Juan_Perez_notion.docx")
    txt = _texto_docx(ruta)
    assert "Marta Ruiz" not in txt
    assert "Le cuesta delegar" not in txt


@pytest.mark.parametrize("incluir,debe_salir", [(True, True), (False, False)])
def test_incluir_fuentes_manda_en_el_word(incluir, debe_salir, tmp_path):
    nombre = f"informe_anual_test_{incluir}.docx"
    sk.guardar_informe_anual_word(
        {"empleado": "Juan Perez", "ca": "Carlos CA", "objetivos": []},
        {"_fuentes": FUENTES},
        nombre_archivo=nombre,
        incluir_fuentes=incluir,
    )
    txt = _texto_docx(str(tmp_path / nombre))
    assert ("Marta Ruiz" in txt) is debe_salir


@pytest.mark.parametrize("incluir,debe_salir", [(True, True), (False, False)])
def test_incluir_fuentes_manda_en_el_html(incluir, debe_salir, tmp_path):
    slug = sk.guardar_informe_anual_html(
        {"empleado": "Juan Perez", "ca": "Carlos CA", "objetivos": []},
        {"_fuentes": FUENTES},
        incluir_fuentes=incluir,
    )
    html = (tmp_path / f"informe_anual_{slug}.html").read_text(encoding="utf-8")
    assert ("Marta Ruiz" in html) is debe_salir


# ── El documento del CA no se le sirve al advisee ────────────────────────────

@pytest.fixture
def como_advisee_con_acceso(as_session, user_session, monkeypatch, tmp_path):
    """El propio evaluado, con el acceso individual que su CA le ha concedido."""
    as_session(user_session)
    monkeypatch.setattr(files_router, "obtener_advisees", lambda *a, **k: [])
    monkeypatch.setattr(files_router, "obtener_ca_de_empleado", lambda ev: "Otra CA")
    monkeypatch.setattr(files_router, "advisee_tiene_acceso_individual", lambda ev, ca: True)
    monkeypatch.setattr(files_router.config, "CARPETA_WEB", str(tmp_path))
    return user_session["persona"]


@pytest.mark.parametrize("nombre", [
    "informe_final_Carlos_CA_1750000000.docx",
    "informe_final_Carlos_CA_1750000000.html",
])
def test_advisee_no_descarga_la_copia_publicada_por_el_ca(client, como_advisee_con_acceso, tmp_path, nombre):
    """La copia con timestamp lleva el anexo de fuentes. El timestamp es un Unix en segundos,
    adivinable a mano, así que el permiso no puede depender de que no lo conozca."""
    (tmp_path / nombre).write_bytes(b"informe con anexo de fuentes")
    r = client.get(f"/api/files/{nombre}", params={"evaluado": como_advisee_con_acceso})
    assert r.status_code == 403
    assert "Solo el CA o un administrador" in r.json()["error"]


def test_advisee_si_descarga_su_informe_final_estable(client, como_advisee_con_acceso, tmp_path):
    """El contrapunto: la versión sin fuentes es justo la que sí debe poder bajarse."""
    (tmp_path / "informe_final_Carlos_CA.docx").write_bytes(b"informe sin anexo")
    r = client.get("/api/files/informe_final_Carlos_CA.docx",
                   params={"evaluado": como_advisee_con_acceso})
    assert r.status_code == 200


# ── 'Ver versión actual del informe' (CA) sí lleva anexo ─────────────────────

def test_la_revision_del_ca_lleva_el_anexo_de_fuentes(tmp_path):
    """Es la razón de ser de esa vista: repasar el informe viendo de dónde sale cada cosa."""
    ea.word_desde_borrador(_borrador(), "revision_informe_Juan_Perez.docx", fuentes=FUENTES)
    txt = _texto_docx(str(tmp_path / "revision_informe_Juan_Perez.docx"))
    assert "FUENTES / EVIDENCIA" in txt.upper()
    assert "Marta Ruiz" in txt


def test_la_revision_se_escribe_en_fichero_propio_y_con_fuentes(client, as_session, user_session,
                                                                monkeypatch, tmp_path):
    """El fichero de la revisión no puede ser el mismo que descarga el advisee: si lo fuera,
    publicar el anexo aquí se lo filtraría por la puerta de al lado."""
    from backend.api.routers import reports as reports_router

    as_session(user_session)
    monkeypatch.setattr(reports_router, "obtener_advisees", lambda *a, **k: ["Juan Perez"])
    monkeypatch.setattr(reports_router, "obtener_informe_estructurado_actual",
                        lambda ev: {"estado": "Final", "contenido": {"empleado": "Juan Perez",
                                                                     "caActual": "Carlos CA",
                                                                     "dimensiones": []}})
    monkeypatch.setattr(reports_router, "mammoth", None)
    recibido = {}

    def _fake(borrador, nombre, fuentes=None):
        recibido["nombre"], recibido["fuentes"] = nombre, fuentes
        return str(tmp_path / nombre)

    monkeypatch.setattr(reports_router.eval_anual_sesion, "word_desde_borrador", _fake)
    monkeypatch.setattr(reports_router.eval_anual_sesion, "fuentes_para_revision", lambda ev: FUENTES)

    r = client.get("/api/informe-final", params={"evaluado": "Juan Perez", "incluir_borrador": 1})
    assert r.status_code == 200
    assert recibido["nombre"].startswith("revision_informe_")
    assert recibido["fuentes"] == FUENTES
