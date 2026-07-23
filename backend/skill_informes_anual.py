"""
Skill: Informe anual IGENERIS
No requiere ninguna base de Notion adicional. Usa las bases existentes:
  - "Evaluaciones - {nombre}"  → evaluaciones mensuales
  - "Opiniones - {nombre}"     → opiniones del CA
  - "Objetivos - {nombre}"     → objetivos por persona (y revela el nombre del CA)
"""

import hashlib
import html as html_lib
import json
import logging
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone

from . import config
from .clients import Document, anthropic_client
from .excepciones import ErrorIA
from .i18n import t
from .ia import CONTACTO, MSG_NO_DISPONIBLE
from .notion_service import (
    excluir_feedback_confidencial,
    listar_bbdd_evaluados,
    obtener_ca_de_empleado,
    obtener_evaluaciones_por_evaluado,
    obtener_opiniones_ca_por_advisee,
    obtener_objetivos_persona,
    obtener_criterios_evaluacion,
    obtener_dimensiones_evaluacion,
    obtener_grupo_empleado,
    obtener_comentarios_personales,
    obtener_barbecho_por_empleado,
    idioma_de_persona,
)
from .project_evals import obtener_evaluaciones_proyecto_por_evaluado
from .evaluaciones_extra import obtener_evaluaciones_extra_por_evaluado
from .utils import normalizar_nombre, slug_archivo


# Etiquetas en ingles para elementos cuya version espanola ya vive en el codigo.
# Se traducen por clave; si no hay traduccion, cae a la etiqueta espanola original.
_MESES_EN = ["January", "February", "March", "April", "May", "June",
             "July", "August", "September", "October", "November", "December"]
_MESES_PT = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
             "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
_DIMS_EN = {
    "gestion_proyecto": "Project management",
    "calidad_tecnica": "Technical quality",
    "trabajo_en_equipo": "Teamwork",
    "comunicacion": "Communication",
    "relacion_cliente": "Client relationship",
    "liderazgo_desarrollo_talento": "Talent development",
    "liderazgo_motivacion": "Motivation",
    "liderazgo_referente": "Role model",
}
_DIMS_PT = {
    "gestion_proyecto": "Gestão de projeto",
    "calidad_tecnica": "Qualidade técnica",
    "trabajo_en_equipo": "Trabalho em equipa",
    "comunicacion": "Comunicação",
    "relacion_cliente": "Relação com o cliente",
    "liderazgo_desarrollo_talento": "Desenvolvimento de talento",
    "liderazgo_motivacion": "Motivação",
    "liderazgo_referente": "Modelo a seguir",
}
_NIVEL_EN = {
    "lider": "Lead",
    "equipo": "Your team members",
    "sin_nivel": "No level specified",
}
_NIVEL_PT = {
    "lider": "Líder",
    "equipo": "Os teus membros da equipa",
    "sin_nivel": "Sem nível especificado",
}

_MESES_POR_IDIOMA = {"en": _MESES_EN, "pt": _MESES_PT}
_DIMS_POR_IDIOMA = {"en": _DIMS_EN, "pt": _DIMS_PT}
_NIVEL_POR_IDIOMA = {"en": _NIVEL_EN, "pt": _NIVEL_PT}


def _mes_label(idx: int, idioma: str) -> str:
    return _MESES_POR_IDIOMA.get(idioma, _MESES_ES)[idx]


# Las dimensiones ya no son una lista fija: salen de Notion y su slug se deriva del
# titulo. Estos son los slugs derivados que corresponden a las claves historicas de
# arriba, para que la traduccion de etiquetas siga encontrandolas.
_SLUG_A_LEGACY = {
    "gestion_del_proyecto":   "gestion_proyecto",
    "relacion_con_el_cliente": "relacion_cliente",
}


def _dim_label(slug: str, etiqueta: str, idioma: str) -> str:
    """Etiqueta de dimension en el idioma dado; conserva el espanol si no hay traduccion."""
    return _DIMS_POR_IDIOMA.get(idioma, {}).get(_SLUG_A_LEGACY.get(slug, slug), etiqueta)


def _nivel_label(clave: str, etiqueta: str, idioma: str) -> str:
    return _NIVEL_POR_IDIOMA.get(idioma, {}).get(clave, etiqueta)


# ── Constantes ────────────────────────────────────────────────────────────────

_REQUIERE_LIDERAZGO = {"sr associate", "manager", "director"}
_CONTENT_W_IN = 9906 / 1440  # ~6.88 pulgadas (A4 márgenes 1.76 cm)

# Fallback de las dimensiones de proyecto: las reales salen de la BD de criterios del
# grupo en Notion (ver dimensiones_informe). Esta lista solo entra si Notion no
# responde o la BD está vacía, para no generar un informe sin apartados.
_DIMS_PROYECTOS = [
    ("gestion_proyecto",  "Gestión del proyecto"),
    ("calidad_tecnica",   "Calidad técnica"),
    ("trabajo_en_equipo", "Trabajo en equipo"),
    ("comunicacion",      "Comunicación"),
    ("relacion_cliente",  "Relación con el cliente"),
]
_DIMS_LIDERAZGO = [
    ("liderazgo_desarrollo_talento", "Desarrollo de Talento"),
    ("liderazgo_motivacion",         "Motivación"),
    ("liderazgo_referente",          "Referente"),
]

_W_DIM  = 3.50
_W_NOTA = 0.60
_W_COM  = _CONTENT_W_IN - _W_DIM - _W_NOTA

_LABELS_NIVEL = [
    ("lider",    "Líder"),
    ("equipo",   "Miembros de tu equipo"),
    ("sin_nivel","Sin nivel especificado"),
]

# Claves de comentarios que son texto plano con bullets. Las dimensiones ya no se pueden
# enumerar aquí (dependen de Notion), así que se reconocen por forma: un dict de
# lider/equipo/sin_nivel es una dimensión, y 'resultado' se queda fuera de ambas a
# propósito, como antes, por ser una síntesis sin citas propias.
_CLAVES_PLANAS = {"contribution_to_firm", "evaluaciones_adicionales", "sin_clasificar"}

# Token de cita por tipo de fuente:
#   E = evaluación mensual · O = opinión CA · P = evaluación de proyecto
#   S = seguimiento personal · B = barbecho · X = evaluación extra (fuera de proyecto)
#   C = aportación del CA en la sesión asistida (lo que sabe y no estaba registrado)
_CITE_RE = re.compile(r"\[([EOPSBXC]\d+)\]")

_MESES_ES = [
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre",
]


def _fmt_deadline(valor) -> str:
    """Formatea una fecha ISO 'YYYY-MM-DD' como 'DD/MM/YYYY'; deja el resto del texto igual."""
    valor = str(valor or "").strip()
    m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", valor)
    return f"{m.group(3)}/{m.group(2)}/{m.group(1)}" if m else valor

def _dims_fijas(pares: list[tuple[str, str]]) -> list[dict]:
    """Convierte una lista fija (clave, etiqueta) al formato de dimensión."""
    return [{"clave": c, "slug": c, "etiqueta": e} for c, e in pares]


def dimensiones_informe(nombre: str, cargo: str, incluir_liderazgo: bool = True,
                        grupo: str = "") -> list[dict]:
    """Dimensiones del informe de esa persona, en orden, como [{clave, slug, etiqueta}].

    Las de proyecto salen de la BD de criterios de SU GRUPO en Notion, así que añadir,
    quitar o reordenar una fila allí cambia los apartados del informe sin tocar código.
    La clave es el id de la página de Notion: sobrevive a que se renombre el criterio,
    que es lo que dejaría huérfano el trabajo ya escrito por el CA.

    Las de liderazgo NO salen de Notion: siguen siendo fijas y condicionadas al cargo.
    """
    try:
        grupo = grupo or _grupo_empleado(nombre, cargo)
        dims = [dict(d) for d in obtener_dimensiones_evaluacion(grupo)]
    except Exception:
        # Notion caído no puede dejar sin informe: se sigue con las dimensiones de siempre.
        logging.exception("[informe] fallo leyendo dimensiones del grupo '%s'", grupo or "?")
        dims = []
    if not dims:
        # Notion caído o BD vacía: mejor el informe de siempre que uno sin apartados.
        logging.warning("[informe] sin dimensiones en Notion para '%s': se usan las fijas", grupo)
        dims = _dims_fijas(_DIMS_PROYECTOS)
    if incluir_liderazgo and any(c in cargo.strip().lower() for c in _REQUIERE_LIDERAZGO):
        dims += _dims_fijas(_DIMS_LIDERAZGO)
    return dims


def huella_dimensiones(dims: list[dict]) -> str:
    """Firma de la estructura del informe. Cambia si se añade, quita, renombra o
    reordena una dimensión, y con ella la huella de caché: así el informe se regenera
    solo cuando de verdad ha cambiado algo, sin gastar IA el resto de las veces."""
    return "|".join(f"{d['clave']}:{d['etiqueta']}" for d in dims)


# ── Criterios DTI por cargo ───────────────────────────────────────────────────

_CRITERIOS_DTI: dict[str, dict[str, list[str]]] = {
    "gestion_proyecto": {
        "analyst": [
            "Priorizar tareas y repartir de forma adecuada los tiempos",
            "Entregar su trabajo a tiempo",
            "Responsabilizarse del buen devenir de sus tareas y subtareas sin necesidad de que se lo recuerden",
            "Es proactivo, detecta necesidades del proyecto y cómo puede aportar valor antes de que alguien se lo diga",
            "Demuestra un compromiso alto hacia un resultado excelente del proyecto",
            "Detecta y avisa de cuellos de botella o posibles problemas intentando aportar soluciones",
            "Muestra disposición y proactividad para encontrar las herramientas que necesita",
            "Demuestra compromiso con las necesidades del proyecto (puntualidad, carga de trabajo, flexibilidad)",
            "Demuestra flexibilidad y motivación hacia la materia del proyecto independientemente de preferencias personales",
        ],
        "associate": [
            "Define y ejecuta con autonomía el plan de trabajo de su área de responsabilidad",
            "Responsabilizarse del proyecto y sus necesidades (desbloquear problemas, establecer reuniones, puntos de seguimiento)",
            "Responsabilizarse de los tiempos del proyecto y de la calidad de los entregables",
            "Identificar las piezas y elementos necesarios para la consecución de un proyecto (herramientas, workshops, sesiones, discusiones internas, con cliente, etc.)",
            "Gestiona adecuadamente y vela por la consecución de todos los elementos necesarios internos",
            "Distribuye adecuadamente las tareas entre los miembros del equipo según cargas de trabajo y perfiles",
            "Vela por mantener un ritmo de trabajo apropiado anticipándose a cuellos de botella o picos de trabajo",
            "Identifica y comunica al responsable del proyecto posibles riesgos y bloqueos",
            "Se focaliza en lo que es más importante (80/20)",
        ],
        "associate sr": [
            "Define el planning de proyecto en profundidad identificando los puntos más complicados",
            "Define el alcance y marco de trabajo del proyecto y lo ajusta de forma continua a la realidad",
            "Identifica nuevas oportunidades para Igeneris que puedan surgir del proyecto (upselling, cross selling)",
            "Es capaz de gestionar un proyecto (estándar y no estándar) entendiendo las necesidades del cliente y ajustando el marco",
            "Se anticipa a posibles riesgos del proyecto y lidera sus posibles planes de contingencia",
        ],
        "manager": [
            "Todo lo de Associate Sr, más:",
            "Ejerce una buena gestión de los tiempos en la organización del proyecto",
            "Sigue una metodología/sello Igeneris, sumada a una base estratégica",
            "Prevé la organización y los posibles riesgos del proyecto, propone un plan de priorización de tareas",
            "Gestiona y resuelve problemas que surgen a lo largo del proyecto",
        ],
    },
    "calidad_tecnica": {
        "analyst": [
            "Se esfuerza y preocupa por entregar su trabajo con máxima calidad",
            "El trabajo que presenta no necesita ser revisado (más de lo necesario) por un tercero",
            "Adquiere y pone en práctica los conocimientos básicos del proyecto (sector, metodología, digitales)",
            "Adquiere un criterio propio sobre la materia del proyecto o tarea",
            "Maneja las herramientas y programas utilizadas en el día a día",
            "Tiene ojo (auto-)crítico para evaluar que la calidad de un trabajo esté conforme con las necesidades de la tarea",
            "Demuestra solvencia en la parte numérica del proyecto si aplica",
        ],
        "associate": [
            "Muestra mediante el ejemplo el nivel de calidad que se ha de cuidar en cada fase del proyecto, sirviendo de guía o referencia",
            "Vela por que el trabajo de los analistas/en prácticas tenga la calidad técnica requerida",
            "Desarrolla la línea de pensamiento y razonamiento numérica necesaria (modelos financieros, magnitudes, economics)",
            "Reta los conceptos numéricos o cualitativos desarrollados para asegurar su rigor",
            "Mantiene el orden en el proyecto — gestión externa, interna y documental, asegurando que la información esté disponible y sea útil",
            "Demuestra madurez en las ideas y tareas en las que trabaja",
            "Aporta un valor fundamental en hipótesis, conclusiones, recomendaciones y presentaciones finales",
        ],
        "associate sr": [
            "Asegura una coherencia a nivel de proyecto en el discurso, el racional, los economics",
            "Identifica de forma rápida las carencias de un proyecto o entregable y sabe subsanarlas eficientemente",
            "Apuesta por ir 'más allá' en los proyectos y traslada al equipo cómo hacerlo",
            "Propone nuevas formas creativas de solucionar problemas",
        ],
        "manager": [
            "Ejerce de referente técnico y estratégico",
            "Se empapa y aterriza el conocimiento sobre la industria en la que se trabaja",
            "Utiliza su experiencia previa en otros proyectos",
            "Garantiza el correcto funcionamiento del equipo y la calidad del entregable a cliente en tiempo y forma",
        ],
    },
    "trabajo_en_equipo": {
        "analyst": [
            "Sabe levantar la mano cuando no tiene capacidad para hacer una tarea",
            "Sus compañeros confían en él porque demuestra un ownership de sus tareas",
            "Se muestra disponible para ayudar a otros compañeros cuando lo necesitan",
            "Contribuye proactivamente al buen clima en el equipo",
            "Acepta las dinámicas de trabajo en equipo y contribuye al buen funcionamiento del equipo",
            "Apoya a sus compañeros en aquellos ámbitos en los que puedan necesitar ayuda",
            "Se preocupa de aprender de sus compañeros y estar al mismo nivel de conocimientos relativos al proyecto",
        ],
        "associate": [
            "Está disponible y accesible para atender a los diferentes miembros de su equipo y guiarlos",
            "Guía al equipo con el ejemplo",
            "Se encarga de que el equipo esté al mismo nivel de información y conocimientos, y da apoyo técnico cuando se necesite",
            "Se asegura que los tiempos dedicados por los analistas/becarios en cada tarea sean los adecuados",
        ],
        "associate sr": [
            "Es un referente para el equipo en cuanto a forma de trabajar y aspiración dentro de Igeneris",
            "Ayuda al equipo a sacar lo mejor de ellos y superar su nivel de calidad, ayudándoles a desarrollarse profesionalmente",
            "Gestiona de forma eficiente la distribución de trabajo del equipo según tiempos y perfiles",
            "Transmite de forma contundente feedback de mejora a los compañeros asegurando que se desarrollen correctamente",
            "Inspira y motiva al equipo a todos los niveles de la pirámide",
            "Está atento a bloqueos y problemas",
        ],
        "manager": [
            "Es capaz de repartir un rol a cada miembro del equipo dentro de las responsabilidades y capacidades reales de cada uno",
            "Es transparente con el equipo",
            "Ejerce con criterio el reparto de tareas en relación con los recursos individuales de cada miembro y el tiempo de inversión",
            "Desarrolla un plan de priorización de tareas donde su equipo pueda entender cuáles son los objetivos y cómo organizarse",
        ],
    },
    "comunicacion": {
        "analyst": [
            "Demuestra una comunicación (oral y escrita) efectiva y asertiva",
            "Muestra una buena comunicación no verbal",
            "Comunica de forma efectiva su criterio al resto del equipo",
            "Demuestra capacidad de razonar sobre su criterio y modificarlo si fuese incorrecto o necesario",
        ],
        "associate": [
            "Comunica de forma efectiva las tareas y prioridades a todos los miembros del equipo",
            "Guía y motiva a los miembros del equipo para sacar lo mejor de ellos y mantener un clima de trabajo positivo",
            "Transmite de forma certera las necesidades del proyecto, especialmente cuando requiere un esfuerzo especial",
            "Sabe construir el storytelling y el racional de una idea, explicársela desde cero a un interlocutor y convencerle de que tiene sentido",
        ],
        "associate sr": [
            "Tiene una alta capacidad de síntesis de los problemas y de exposición tanto internamente como hacia cliente",
            "Comunica de forma clara a todos los niveles de la organización del cliente adaptando el discurso y contenido a cada auditorio",
            "Argumenta con seguridad y convincentemente, siendo capaz de reaccionar a argumentaciones del cliente",
        ],
        "manager": [
            "Transparencia en la comunicación a lo largo del proyecto para que el equipo esté alineado con el cliente/proyecto",
            "Sabe dar una comunicación asertiva al equipo",
            "Sabe comunicar al cliente adaptando el discurso dependiendo de las necesidades del proyecto y de las reacciones potenciales del cliente",
        ],
    },
    "relacion_cliente": {
        "analyst": [
            "Participa en reuniones con clientes",
            "Entiende las dinámicas con el cliente y el trato que se le debe dar",
        ],
        "associate": [
            "Define y prepara las sesiones de trabajo con el cliente",
            "Logra confianza y credibilidad con los niveles del cliente con los que le corresponde relacionarse, transmitiendo seguridad y profesionalidad",
            "Lidera sesiones de trabajo con el cliente de forma asistida por alguien con más seniority",
            "Lidera sesiones de trabajo con el cliente de forma autónoma",
        ],
        "associate sr": [
            "Crea un vínculo con el cliente y es capaz de entender sus necesidades para con el proyecto",
            "Lidera los workshops y sesiones de trabajo con el cliente más complicados / coordina y supervisa que las reuniones estén bien pensadas y ejecutadas",
            "Es un referente para el cliente en todos los aspectos que abarca el proyecto e incluso más allá del alcance del mismo",
        ],
        "manager": [
            "Mantiene una buena comunicación con el cliente",
            "Sabe preguntar al cliente qué necesita y cuáles son sus expectativas para no ir apagando fuegos posteriormente",
            "Conduce eficazmente las expectativas del cliente, contribuyendo a la satisfacción con el resultado del proyecto",
        ],
    },
}

_ETIQUETAS_DIM = {
    "gestion_proyecto": "Gestión del proyecto",
    "calidad_tecnica": "Calidad técnica",
    "trabajo_en_equipo": "Trabajo en equipo",
    "comunicacion": "Comunicación",
    "relacion_cliente": "Relación con el cliente",
}

# Escala canonica de niveles de criterios, de menor a mayor. Coincide con los
# nombres de columna de las BD "Criterios de evaluaciones" de Notion (en espanol).
_ORDEN_CARGO = ["trainee", "analista", "asociado", "asociado_sr", "manager"]

# Normaliza cualquier etiqueta de nivel (columna de Notion en ES, clave inglesa
# del fallback hardcodeado, etc.) a la clave canonica de _ORDEN_CARGO.
_NIVEL_ALIAS = {
    "trainee": "trainee", "becario": "trainee", "en practicas": "trainee", "intern": "trainee",
    "analista": "analista", "analyst": "analista",
    "asociado": "asociado", "associate": "asociado",
    "asociado sr": "asociado_sr", "asociado senior": "asociado_sr",
    "associate sr": "asociado_sr", "sr associate": "asociado_sr", "senior associate": "asociado_sr",
    "manager": "manager",
}


def _norm_txt(s: str) -> str:
    return " ".join((s or "").strip().lower().replace(".", " ").split())


def _nivel_canonico(etiqueta: str) -> str | None:
    """Etiqueta de nivel (columna de Notion o clave del fallback) -> clave canonica."""
    return _NIVEL_ALIAS.get(_norm_txt(etiqueta))


def _nivel_cargo(cargo: str) -> str | None:
    """Cargo del empleado (en ingles en la Lista de empleados) -> nivel canonico de
    criterios. Los cargos por encima de Manager (Sr. Manager, Director, Partner,
    Lead/Director de Palantir) usan los criterios de 'Manager'."""
    c = _norm_txt(cargo)
    if not c:
        return None
    # Palantir: titulos de ingenieria.
    if "palantir" in c or "engineer" in c:
        # Ingeniero Jr. = Analista · Ingeniero = Asociado · Ingeniero Sr. = Asociado Sr
        # Lead / Director / Partner = Manager. Trainee se mantiene por si existe.
        if "lead" in c or "director" in c or "partner" in c:
            return "manager"
        if "trainee" in c or "becario" in c or "intern" in c:
            return "trainee"
        if "jr" in c or "junior" in c:
            return "analista"
        if "sr" in c or "senior" in c:
            return "asociado_sr"
        if "engineer" in c:
            return "asociado"
        return None
    # Negocio / general.
    if "partner" in c or "director" in c or "manager" in c:
        return "manager"  # incluye Sr. Manager y Manager
    if "trainee" in c or "becario" in c or "intern" in c:
        return "trainee"
    if "associate" in c or "asociado" in c:
        return "asociado_sr" if ("sr" in c or "senior" in c) else "asociado"
    if "analyst" in c or "analista" in c:
        return "analista"
    return None


def _grupo_por_cargo(cargo: str) -> str:
    """Fallback: infiere el grupo del TEXTO del cargo (poco fiable). Usa _grupo_empleado si tienes el nombre."""
    c = cargo.lower()
    if "palantir" in c:
        return "Palantir"
    if "head" in c:
        return "MiddleOffice"
    return "Negocio"


def _grupo_empleado(nombre: str, cargo: str) -> str:
    """Grupo real (Negocio/Palantir/MiddleOffice) desde la columna Área de Notion.
    Si no consta, cae a inferirlo del cargo."""
    if nombre:
        try:
            grupo = obtener_grupo_empleado(nombre)
            if grupo:
                return grupo
        except Exception:
            logging.exception("No se pudo leer el grupo de '%s' desde Notion", nombre)
    return _grupo_por_cargo(cargo)


def _criterios_para_prompt(cargo: str, idioma: str = "es", nombre: str = "",
                           dims: list[dict] | None = None) -> str:
    """Bloque de criterios para el prompt.

    Con `dims`, cada bloque se rotula con la clave EXACTA que Claude tiene que devolver,
    en vez de con el título del criterio. Sin eso, el modelo recibía las claves por un
    lado ("Dimensiones requeridas: ...") y los criterios por otro, rotulados distinto, y
    acababa agrupando la evidencia por el nombre de la dimensión en vez de por lo que sus
    criterios describen: un criterio vacío o renombrado seguía llevándose las mismas
    evaluaciones de siempre.
    """
    nivel = _nivel_cargo(cargo)
    grupo = _grupo_empleado(nombre, cargo)

    try:
        criterios_notion = obtener_criterios_evaluacion(grupo, idioma)
    except Exception:
        criterios_notion = {}

    def _por_canon(niveles_dict):
        d = {}
        for label, crits in niveles_dict.items():
            canon = _nivel_canonico(label)
            if canon and crits:
                d[canon] = (label, crits)
        return d

    idx = _ORDEN_CARGO.index(nivel) if nivel in _ORDEN_CARGO else -1

    def _lineas(niveles_dict):
        por_canon = _por_canon(niveles_dict)
        presentes = [c for c in _ORDEN_CARGO if c in por_canon]
        sel = presentes[max(0, idx - 1):] if idx >= 0 else presentes
        return [f"  [{por_canon[c][0]}]: " + " / ".join(por_canon[c][1]) for c in sel]

    if criterios_notion:
        bloques = [f"Lo que se espera de un {cargo} en {grupo} y niveles superiores como referencia:"]
        # Rotular por la clave que se le pide devolver, para que no haya que adivinar qué
        # bloque corresponde a qué dimensión.
        etiqueta_a_slug = {d["etiqueta"]: d["slug"] for d in (dims or [])}
        for dim_label, niveles_dict in criterios_notion.items():
            lineas = _lineas(niveles_dict)
            if not lineas:
                continue
            slug = etiqueta_a_slug.get(dim_label)
            cabecera = f'"{slug}" ({dim_label})' if slug else dim_label
            bloques.append(f"\n{cabecera}:\n" + "\n".join(lineas))
        # Una dimensión sin criterios en Notion no debe llevarse evidencia por inercia.
        for d in (dims or []):
            if d["etiqueta"] not in criterios_notion:
                bloques.append(f'\n"{d["slug"]}" ({d["etiqueta"]}):\n  (sin criterios definidos en Notion)')
        return "\n".join(bloques)

    # Fallback al diccionario hardcodeado (solo Negocio)
    if not nivel:
        return ""
    bloques = [f"Lo que se espera de un {cargo} y niveles superiores como referencia:"]
    for dim_key, dim_label in _ETIQUETAS_DIM.items():
        lineas = _lineas(_CRITERIOS_DTI.get(dim_key, {}))
        if lineas:
            bloques.append(f"\n{dim_label}:\n" + "\n".join(lineas))
    return "\n".join(bloques)


# ── Lista de empleados ────────────────────────────────────────────────────────

def obtener_empleados_evaluacion_anual() -> list[str]:
    """
    Devuelve los empleados que tienen base "Evaluaciones - {nombre}" en Notion.
    Son los candidatos válidos para generar el informe anual.
    """
    try:
        bases = listar_bbdd_evaluados()
        return sorted(b["evaluado"] for b in bases if b.get("evaluado"))
    except Exception:
        logging.exception("Error listando empleados para informe anual.")
        return []


# ── Recopilar datos del empleado ──────────────────────────────────────────────

def _resultado(futuro, clave: str, nombre: str, por_defecto=None):
    """Resultado de una lectura, o el valor por defecto si falló.

    Cada lectura falla por su cuenta: si una base no existe o Notion la rechaza, esa
    clave se queda vacía y las demás siguen, igual que hacía el try/except de cada bloque.
    """
    try:
        return futuro.result()
    except Exception:
        logging.warning("No se pudo leer '%s' de %s.", clave, nombre)
        return [] if por_defecto is None else por_defecto


def obtener_datos_empleado_anual(nombre: str) -> dict:
    """
    Recopila toda la información disponible en Notion sobre el empleado:
      - evaluaciones mensuales (desde "Evaluaciones - {nombre}")
      - opiniones del CA (desde "Opiniones - {nombre}")
      - objetivos (desde "Objetivos empleados"), de donde también se extrae el CA

    Todas las lecturas van a la vez. Eran 8 llamadas encadenadas (~30s) en las que cada
    una se pasaba el rato esperando a la red de Notion sin hacer nada; lanzadas juntas
    tardan lo que la más lenta. Medido sobre un empleado real: 30s -> 8s.
    """
    # Un worker por tarea, y uno de más para las opiniones: esa se queda esperando al
    # resultado de otra, y si no tuviera worker propio podría quedarse sin sitio y
    # bloquear el pool entero.
    _t_notion = time.time()
    with ThreadPoolExecutor(max_workers=8) as pool:
        # Independientes entre sí → todas a la vez.
        # `ca_lista` se pide aunque casi siempre sobre (el CA suele venir en los
        # objetivos): aquí sale gratis, y pedirlo después sería otra ida y vuelta.
        fut = {
            "evaluaciones": pool.submit(lambda: excluir_feedback_confidencial(obtener_evaluaciones_por_evaluado(nombre))),
            "objetivos": pool.submit(obtener_objetivos_persona, nombre),
            "evals_proyecto": pool.submit(obtener_evaluaciones_proyecto_por_evaluado, nombre),
            "seguimiento": pool.submit(obtener_comentarios_personales, nombre),
            "barbecho": pool.submit(obtener_barbecho_por_empleado, nombre),
            "evals_extra": pool.submit(obtener_evaluaciones_extra_por_evaluado, nombre),
            "ca_lista": pool.submit(lambda: obtener_ca_de_empleado(nombre) or ""),
        }

        def _ca_y_opiniones():
            """Las opiniones necesitan saber quién es el CA, así que dependen de otra
            lectura. En vez de esperar a que termine todo el bloque, arranca en cuanto
            se sabe el CA (~4s) y se solapa con el resto: es la lectura más lenta (~12s)
            y encadenarla al final añadía su tiempo entero al total."""
            objetivos = _resultado(fut["objetivos"], "objetivos", nombre)
            ca = (objetivos[0].get("ca", "") if objetivos else "") or _resultado(fut["ca_lista"], "ca_lista", nombre, "")
            try:
                return ca, obtener_opiniones_ca_por_advisee(ca, nombre)
            except Exception:
                logging.warning("No se encontraron opiniones del CA para %s.", nombre)
                return ca, []

        fut_opiniones = pool.submit(_ca_y_opiniones)

        evaluaciones = _resultado(fut["evaluaciones"], "evaluaciones", nombre)
        objetivos = _resultado(fut["objetivos"], "objetivos", nombre)
        evals_proyecto = _resultado(fut["evals_proyecto"], "evals_proyecto", nombre)
        seguimiento = _resultado(fut["seguimiento"], "seguimiento", nombre)
        barbecho = _resultado(fut["barbecho"], "barbecho", nombre)
        evals_extra = _resultado(fut["evals_extra"], "evals_extra", nombre)
        ca_nombre, opiniones = _resultado(fut_opiniones, "opiniones_ca", nombre, ("", []))

    logging.warning("[perf] lectura de todo Notion para %s: %.1fs", nombre, time.time() - _t_notion)
    return {
        "empleado": nombre,
        "ca": ca_nombre,
        "opiniones_ca": opiniones,
        "evaluaciones": evaluaciones,
        "evals_proyecto": evals_proyecto,
        "seguimiento": seguimiento,
        "barbecho": barbecho,
        "evals_extra": evals_extra,
        "objetivos": objetivos,
    }


# ── Claude: interpretación ────────────────────────────────────────────────────

_MES_ABBR = ["ene", "feb", "mar", "abr", "may", "jun", "jul", "ago", "sep", "oct", "nov", "dic"]


def _mes_tag(fecha: str) -> str:
    """'2025-03-15' -> 'mar 25'. Para que Claude valore la evolución temporal."""
    try:
        return f"{_MES_ABBR[int(fecha[5:7]) - 1]} {fecha[2:4]}"
    except Exception:
        return "s/f"


def _label(partes: list) -> str:
    """Une las partes no vacías de una etiqueta de fuente con ' · '.

    Las partes son opcionales porque en el flujo asistido llegan redactadas (sin nivel ni
    tipo, ver eval_anual_sesion._redactar_emp_data): un join fijo dejaba separadores
    huérfanos ('Proyecto Alfa ·  · 2026-07-08').
    """
    return " · ".join(str(p) for p in partes if p)


def _formatear_contexto(emp_data: dict) -> tuple[str, dict]:
    """Construye el texto que se pasa a Claude y el mapa de fuentes citables.

    Devuelve (texto, fuentes) donde fuentes mapea cada id de cita
    (``E3``, ``O1``, ``P2``, ``S1``, ``B1``) a ``{"url", "tipo", "label", "texto"}``.
    Cada línea va prefijada con su id y con la etiqueta de mes para que Claude
    cite y, además, valore la evolución a lo largo del año.
    """
    bloques = []
    fuentes: dict[str, dict] = {}

    # ── Opiniones del CA (notas del CA + resúmenes del chatbot) ── [O#] ──────
    opiniones = sorted(emp_data.get("opiniones_ca", []), key=lambda x: (x.get("fecha") or ""))
    if opiniones:
        bloques.append("=== OPINIONES DEL CA (orden cronológico) ===")
        for i, op in enumerate(opiniones, 1):
            fecha = (op.get("fecha") or "")[:10] or "Sin fecha"
            partes = []
            if op.get("resumen_advisee"):
                partes.append(f"Resumen (chatbot): {op['resumen_advisee']}")
            if op.get("opinion"):
                partes.append(f"Nota del CA: {op['opinion']}")
            if not partes:
                continue
            cid = f"O{i}"
            fuentes[cid] = {
                "url": op.get("url", ""), "tipo": "opinion", "fecha": (op.get("fecha") or "")[:10],
                "label": f"Opinión CA · {fecha}", "texto": " | ".join(partes),
            }
            bloques.append(f"[{cid}] [{_mes_tag(fecha)}] " + " | ".join(partes))

    # ── Evaluaciones mensuales, agrupadas por jerarquía y cronológicas ── [E#] ─
    evaluaciones = emp_data.get("evaluaciones", [])
    if evaluaciones:
        _GRUPOS = [
            ("lider",     "EVALUACIONES MENSUALES DEL LÍDER"),
            ("equipo",    "EVALUACIONES MENSUALES DE MIEMBROS DEL EQUIPO (iguales y subordinados)"),
            ("sin_nivel", "EVALUACIONES MENSUALES SIN NIVEL ESPECIFICADO"),
        ]
        _ETIQUETA_REL = {"lider": "líder", "equipo": "equipo", "sin_nivel": "sin nivel"}
        por_rel: dict = {"lider": [], "equipo": [], "sin_nivel": []}
        for ev in evaluaciones:
            rel = ev.get("relacion", "")
            if rel == "superior":
                por_rel["lider"].append(ev)
            elif rel in ("igual", "inferior"):
                por_rel["equipo"].append(ev)
            else:
                por_rel["sin_nivel"].append(ev)
        n = 0
        for rel_key, encabezado in _GRUPOS:
            evs = sorted(por_rel.get(rel_key, []), key=lambda x: (x.get("fecha") or ""))
            if not evs:
                continue
            bloques.append(f"\n=== {encabezado} (orden cronológico) ===")
            for ev in evs:
                n += 1
                cid = f"E{n}"
                proyecto  = ev.get("proyecto") or "Sin proyecto"
                anonimo   = bool(ev.get("anonimizado"))
                evaluador = "" if anonimo else (ev.get("persona_que_evalua") or ev.get("nombre") or "Desconocido")
                fecha     = (ev.get("fecha") or "")[:10]
                q1, q2 = ev.get("q1", ""), ev.get("q2", "")
                fuentes[cid] = {
                    "url": ev.get("url", ""), "tipo": "evaluacion", "fecha": fecha,
                    "label": _label([proyecto, None if anonimo else _ETIQUETA_REL[rel_key], fecha]),
                    "evaluador": evaluador,
                    "texto": f"Valoración: {q1}/5 | Ejemplo: {q2}",
                }
                bloques.append(
                    f"[{cid}] [{_mes_tag(fecha)}] Proyecto: {proyecto} | "
                    + (f"Evaluador: {evaluador} | " if evaluador else "")
                    + f"Valoración: {q1}/5 | Ejemplo: {q2}"
                )

    # ── Evaluaciones de proyecto (por proyecto y cronológicas) ── [P#] ───────
    evals_proy = sorted(emp_data.get("evals_proyecto", []), key=lambda x: (x.get("fecha") or ""))
    if evals_proy:
        bloques.append("\n=== EVALUACIONES DE PROYECTO (orden cronológico) ===")
        for i, pe in enumerate(evals_proy, 1):
            cid = f"P{i}"
            proyecto  = pe.get("proyecto") or "Sin proyecto"
            anonimo   = bool(pe.get("anonimizado"))
            evaluador = "" if anonimo else (pe.get("evaluador") or "Desconocido")
            tipo      = "" if anonimo else (pe.get("tipo") or "")
            fecha     = (pe.get("fecha") or "")[:10]
            respuestas = pe.get("respuestas") or ""
            fuentes[cid] = {
                "url": pe.get("url", ""), "tipo": "proyecto", "fecha": fecha,
                "label": _label([proyecto, tipo, fecha]),
                "evaluador": evaluador,
                "texto": respuestas,
            }
            bloques.append(
                f"[{cid}] [{_mes_tag(fecha)}] Proyecto: {proyecto} | "
                + (f"Evaluador: {evaluador} | " if evaluador else "")
                + (f"Tipo: {tipo} | " if tipo else "")
                + f"Respuestas: {respuestas}"
            )

    # ── Seguimiento personal ── [S#] ─────────────────────────────────────────
    seguimiento = sorted(emp_data.get("seguimiento", []), key=lambda x: (x.get("fecha") or ""))
    if seguimiento:
        bloques.append("\n=== SEGUIMIENTO PERSONAL (orden cronológico) ===")
        for i, sg in enumerate(seguimiento, 1):
            cid = f"S{i}"
            fecha = (sg.get("fecha") or "")[:10]
            autor = sg.get("autor") or ""
            comentario = sg.get("comentario") or ""
            fuentes[cid] = {
                "url": sg.get("url", ""), "tipo": "seguimiento", "fecha": fecha,
                "label": f"Seguimiento personal · {fecha}", "evaluador": autor,
                "texto": comentario,
            }
            bloques.append(f"[{cid}] [{_mes_tag(fecha)}] Seguimiento ({autor}): {comentario}")

    # ── Barbecho (labores sin proyecto → contribución a la firma) ── [B#] ────
    barbecho = sorted(emp_data.get("barbecho", []), key=lambda x: (x.get("fecha") or ""))
    if barbecho:
        bloques.append("\n=== BARBECHO — labores en periodo sin proyecto (orden cronológico) ===")
        for i, bb in enumerate(barbecho, 1):
            cid = f"B{i}"
            fecha = (bb.get("fecha") or "")[:10]
            area = bb.get("area") or ""
            labores = bb.get("labores") or ""
            fuentes[cid] = {
                "url": bb.get("url", ""), "tipo": "barbecho", "fecha": fecha,
                "label": f"Barbecho{f' ({area})' if area else ''} · {fecha}", "texto": labores,
            }
            bloques.append(f"[{cid}] [{_mes_tag(fecha)}] Barbecho ({area}): {labores}")

    # ── Evaluaciones extra (fuera de proyecto) ── [X#] ───────────────────────
    evals_extra = sorted(emp_data.get("evals_extra", []), key=lambda x: (x.get("fecha") or ""))
    if evals_extra:
        bloques.append("\n=== EVALUACIONES EXTRA (fuera de proyecto, orden cronológico) ===")
        for i, ee in enumerate(evals_extra, 1):
            cid = f"X{i}"
            evaluador = ee.get("evaluador") or "Desconocido"
            contexto = ee.get("contexto") or "Sin contexto"
            fecha = (ee.get("fecha") or "")[:10]
            respuestas = ee.get("respuestas") or ""
            fuentes[cid] = {
                "url": ee.get("url", ""), "tipo": "extra", "fecha": fecha,
                "label": f"{contexto} · {fecha}".strip(" ·"),
                "evaluador": evaluador,
                "texto": respuestas,
            }
            bloques.append(
                f"[{cid}] [{_mes_tag(fecha)}] Contexto: {contexto} | Evaluador: {evaluador} | "
                f"Respuestas: {respuestas}"
            )

    # ── Aportaciones del CA en la sesión asistida ── [C#] ────────────────────
    # No salen de Notion: las aporta el CA hablando con la IA y solo se registran cuando
    # concreta cuándo, dónde y qué pasó (ver eval_anual_sesion._registrar_aportaciones).
    # Una vez registradas son fuente citable como cualquier otra: llevan autor y fecha, y
    # aparecen en el anexo del informe con su texto literal.
    for ap in emp_data.get("aportaciones_ca", []):
        cid = ap.get("cid")
        if not cid:
            continue
        fecha = (ap.get("fecha") or "")[:10]
        fuentes[cid] = {
            "url": "", "tipo": "aportacion_ca", "fecha": fecha,
            "label": _label(["Aportación del CA", ap.get("proyecto"), fecha]),
            "evaluador": ap.get("autor") or "",
            "texto": ap.get("texto") or "",
        }
        bloques.append(
            f"\n[{cid}] [{_mes_tag(fecha)}] Aportación del CA"
            + (f" | Proyecto: {ap['proyecto']}" if ap.get("proyecto") else "")
            + f" | {ap.get('texto') or ''}"
        )

    texto = "\n".join(bloques) if bloques else "(Sin datos de evaluación disponibles)"
    return texto, fuentes


# Lo que escribe Claude cuando una dimensión no tiene evidencia. Va sin cita a propósito,
# así que hay que reconocerlo para no confundirlo con una afirmación sin respaldo.
_SIN_INFO = ("sin informacion suficiente", "sem informacao suficiente", "not enough information")


def _es_sin_informacion(linea: str) -> bool:
    # Plegando tildes: _norm_txt no las quita, así que "información" nunca casaría.
    return normalizar_nombre(_norm_txt(linea)).rstrip(". ") in _SIN_INFO


def _filtrar_bullets_citados(texto: str, fuentes: dict, descartados: list) -> str:
    """Devuelve solo los bullets que tienen al menos una cita válida.

    - Bullet sin ninguna cita -> se descarta (no hay forma de comprobar su origen).
    - Citas a ids inexistentes -> se eliminan del texto.
    Esto convierte cualquier invención en algo estructural: si Claude no puede
    señalar de qué evaluación sale una afirmación, la afirmación no aparece.
    """
    lineas_ok = []
    for linea in (texto or "").splitlines():
        bruta = linea.strip(" •-–\t")
        if not bruta:
            continue
        if _es_sin_informacion(bruta):
            # El prompt pide escribir esto SIN cita cuando una dimensión no tiene
            # evidencia. Descartarlo por no citar dejaba la dimensión vacía y borraba
            # justo la señal de que ahí no hay nada que decir.
            lineas_ok.append(bruta)
            continue
        ids = _CITE_RE.findall(bruta)
        validos = [i for i in ids if i in fuentes]
        if not validos:
            descartados.append(bruta)
            continue
        # Elimina citas a ids inexistentes, conserva las válidas
        invalidos = set(ids) - set(validos)
        for inv in invalidos:
            bruta = bruta.replace(f"[{inv}]", "")
        lineas_ok.append(re.sub(r"\s{2,}", " ", bruta).strip())
    return "\n".join(lineas_ok)


def _validar_citas(comentarios: dict, fuentes: dict) -> dict:
    """Aplica el filtro de citas a todas las dimensiones y registra lo descartado."""
    descartados: list[str] = []
    for clave, valor in list(comentarios.items()):
        if clave.startswith("_"):
            continue
        if isinstance(valor, dict):
            comentarios[clave] = {
                nivel: _filtrar_bullets_citados(txt, fuentes, descartados)
                for nivel, txt in valor.items()
            }
        elif clave in _CLAVES_PLANAS and isinstance(valor, str):
            comentarios[clave] = _filtrar_bullets_citados(valor, fuentes, descartados)
    if descartados:
        logging.warning(
            "[informe] %d bullet(s) descartados por no citar ninguna fuente válida: %s",
            len(descartados), descartados,
        )
    comentarios["_bullets_descartados"] = descartados
    return comentarios


def _remapear_slugs_a_claves(comentarios: dict, dims: list[dict]) -> dict:
    """Traduce las claves que devuelve Claude (slugs) a las claves persistidas.

    Se conserva el orden de `dims` para que el informe salga en el orden de Notion, y las
    claves que no son dimensiones (contribution_to_firm, resultado, los metadatos con _)
    pasan tal cual.
    """
    por_slug = {d["slug"]: d["clave"] for d in dims}
    salida: dict = {}
    for d in dims:
        if d["slug"] in comentarios:
            salida[d["clave"]] = comentarios[d["slug"]]
    for clave, valor in comentarios.items():
        if clave in por_slug:
            continue
        salida[clave] = valor
    return salida


def _recolectar_afirmaciones(comentarios: dict) -> list[dict]:
    """Extrae cada bullet con sus citas para auditarlo. No incluye 'resultado' (síntesis)."""
    afirmaciones = []
    for clave, valor in comentarios.items():
        if clave.startswith("_"):
            continue
        if isinstance(valor, dict):
            textos = [(nivel, t) for nivel, t in valor.items()]
        elif clave in _CLAVES_PLANAS and isinstance(valor, str):
            textos = [("", valor)]
        else:
            continue
        for _nivel, bloque in textos:
            for linea in (bloque or "").splitlines():
                bruta = linea.strip(" •-–\t")
                if not bruta:
                    continue
                ids = _CITE_RE.findall(bruta)
                if ids:
                    afirmaciones.append({"clave": clave, "texto": bruta, "citas": ids})
    return afirmaciones


def _verificar_soporte(comentarios: dict, fuentes: dict) -> list[dict]:
    """Segunda pasada (auditor): marca afirmaciones NO respaldadas por su cita.

    Política: avisar, no borrar. Devuelve una lista de avisos; el CA decide.
    Si la llamada falla, devuelve [] (el informe sigue siendo válido).
    """
    if not anthropic_client:
        return []
    afirmaciones = _recolectar_afirmaciones(comentarios)
    if not afirmaciones:
        return []

    # Solo las fuentes realmente citadas, para acotar el contexto
    citadas = sorted({cid for a in afirmaciones for cid in a["citas"]})
    bloque_fuentes = "\n".join(
        f"[{cid}] {fuentes[cid].get('texto', '')}" for cid in citadas if cid in fuentes
    )
    bloque_afirmaciones = "\n".join(
        f"{i}. \"{a['texto']}\"" for i, a in enumerate(afirmaciones)
    )

    system = (
        "Eres un auditor de calidad de informes de RRHH. Recibes (A) el texto literal de las "
        "fuentes y (B) una lista numerada de afirmaciones, cada una con las citas [E#]/[O#] de "
        "las que dice provenir. Tu ÚNICA tarea es marcar las afirmaciones que NO estén "
        "respaldadas por el texto literal de su(s) cita(s): inventadas, exageradas, o que "
        "afirman más de lo que la fuente dice. Sé estricto pero justo: una reformulación fiel "
        "SÍ está respaldada. Devuelve ÚNICAMENTE un JSON válido: "
        '{"no_soportadas": [{"i": <numero de la afirmacion>, "motivo": "<breve>"}]}. '
        "Si todas están respaldadas, devuelve {\"no_soportadas\": []}."
        + config.INSTRUCCION_ANTIINYECCION
    )
    try:
        respuesta = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1500,
            temperature=0,
            system=system,
            messages=[{
                "role": "user",
                "content": f"=== FUENTES ===\n{bloque_fuentes}\n\n=== AFIRMACIONES ===\n{bloque_afirmaciones}",
            }],
        )
        texto = "".join(b.text for b in respuesta.content if b.type == "text").strip()
        data = _extraer_json_objeto(texto)
        if data is None:
            logging.error("[informe] Verificación no parseable; se omite. Texto: %r", texto[:800])
            return []
    except Exception:
        logging.exception("[informe] Falló la pasada de verificación; se omite.")
        return []

    avisos = []
    for item in data.get("no_soportadas", []):
        try:
            idx = int(item.get("i"))
        except (TypeError, ValueError):
            continue
        if 0 <= idx < len(afirmaciones):
            a = afirmaciones[idx]
            avisos.append({
                "afirmacion": a["texto"],
                "clave": a["clave"],
                "citas": a["citas"],
                "motivo": (item.get("motivo") or "").strip(),
            })
    if avisos:
        logging.warning("[informe] Verificador marcó %d afirmación(es) no respaldadas.", len(avisos))
    return avisos


def _bloque_equilibrado(t: str, inicio: int) -> str | None:
    """Recorta desde la '{' de `inicio` hasta su '}' pareja. None si no cierra.

    Contar llaves a pelo no vale: las que aparecen dentro de las cadenas del JSON
    (y las comillas escapadas) descuadran el recuento, así que hay que seguir el estado.
    """
    profundidad, en_cadena, escapado = 0, False, False
    for i in range(inicio, len(t)):
        c = t[i]
        if en_cadena:
            if escapado:
                escapado = False
            elif c == "\\":
                escapado = True
            elif c == '"':
                en_cadena = False
        elif c == '"':
            en_cadena = True
        elif c == "{":
            profundidad += 1
        elif c == "}":
            profundidad -= 1
            if profundidad == 0:
                return t[inicio:i + 1]
    return None


def _extraer_json_objeto(texto: str) -> dict | None:
    """Devuelve el objeto JSON de la respuesta del modelo, o None si no hay ninguno.

    Aunque se le pida solo JSON, el modelo a veces antepone su razonamiento ("I need to
    analyze all sources...") o añade comentarios después. No basta con recortar por la
    primera '{': esa prosa puede traer llaves sueltas o un JSON de ejemplo. Así que se
    prueban todos los candidatos y se elige el objeto válido más largo, que es el informe
    (cualquier ejemplo suelto del razonamiento es mucho más pequeño).
    """
    t = texto.strip()
    if t.startswith("```"):
        t = t.split("```", 2)[1]
        if t.startswith("json"):
            t = t[4:]
        t = t.rsplit("```", 1)[0].strip()
    mejor, mejor_len = None, 0
    i = t.find("{")
    while i >= 0:
        bloque = _bloque_equilibrado(t, i)
        if bloque:
            try:
                datos = json.loads(bloque)
            except json.JSONDecodeError:
                datos = None
            if isinstance(datos, dict) and len(bloque) > mejor_len:
                mejor, mejor_len = datos, len(bloque)
                # El resto de llaves de este bloque son suyas: no hay nada mejor dentro.
                i = t.find("{", i + len(bloque))
                continue
        i = t.find("{", i + 1)
    return mejor


def interpretar_evaluaciones_anual(emp_data: dict, cargo: str = "", criterios: str | None = None,
                                   idioma: str = "es", dims: list[dict] | None = None) -> dict:
    """
    Llama a Claude con el contexto de evaluaciones y opiniones.
    Devuelve un dict con bullets y notas por dimensión, tecleado por la clave de cada una.

    `criterios`: texto de criterios ya renderizado. Si es None se obtiene de Notion.
    Pasarlo evita una segunda lectura de Notion cuando ya se computó para la huella de caché.
    `dims`: dimensiones ya resueltas, por el mismo motivo.
    """
    if not anthropic_client:
        raise ErrorIA(MSG_NO_DISPONIBLE, "ia_no_configurada", definitivo=True)

    cargo_lower = cargo.strip().lower()
    requiere_liderazgo = any(c in cargo_lower for c in _REQUIERE_LIDERAZGO)

    if dims is None:
        dims = dimensiones_informe(emp_data.get("empleado", ""), cargo)
    # A Claude se le piden los slugs, no las claves: la clave es un id de página de Notion
    # y meter UUIDs en el prompt lo haría ilegible y fácil de equivocar. Se remapean al
    # volver, que es donde importa que la clave sea estable.
    dims_lista = ", ".join(f'"{d["slug"]}"' for d in dims)

    criterios_bloque = (_criterios_para_prompt(cargo, idioma, emp_data.get("empleado", ""), dims=dims)
                        if criterios is None else criterios)
    criterios_section = (
        "\n\nCRITERIOS DE CADA DIMENSIÓN (definen QUÉ evalúa cada una, y calibran el nivel "
        "exigible según el cargo):\n" + criterios_bloque +
        "\n\nQUÉ VA EN CADA DIMENSIÓN (obligatorio):\n"
        "Esto aplica SOLO a las dimensiones listadas arriba con sus criterios. Los criterios son "
        "su DEFINICIÓN, no un adorno: coloca cada evidencia en la dimensión cuyos criterios "
        "describan lo que esa evidencia cuenta. No te guíes por el nombre de la dimensión, "
        "guíate por sus criterios. Si los criterios de una dimensión cambian, cambia también qué "
        "evidencia le corresponde.\n"
        "Si una dimensión no tiene criterios definidos, o ninguna evidencia encaja con los que "
        "tiene, escribe exactamente 'Sin información suficiente'. NO le asignes evidencia solo "
        "porque su nombre suene relacionado o porque antes fuera de otra manera.\n"
        "EXCEPCIÓN — 'contribution_to_firm', 'evaluaciones_adicionales' y 'resultado' NO tienen "
        "criterios y NO se rigen por esta regla: siguen guiándose por su significado de siempre "
        "(contribution_to_firm = contribución a la firma, sobre todo barbecho [B#] y seguimiento "
        "personal marcado como CTTF)."
        if criterios_bloque else ""
    )

    contexto, fuentes = _formatear_contexto(emp_data)

    system = (
        "Eres el director de RRHH de IGENERIS. A partir de TODAS las fuentes del empleado, "
        "genera el contenido del informe anual de evaluación. "
        "Ten en cuenta los criterios DTI: lo que es suficiente para un Analyst puede ser lo mínimo esperado para un Manager.\n\n"
        "FUENTES (cada línea lleva su etiqueta y el mes entre corchetes):\n"
        "  [O#] opinión del CA (sus notas + resúmenes del chatbot)\n"
        "  [E#] evaluación mensual (con jerarquía líder/equipo)\n"
        "  [P#] evaluación de proyecto\n"
        "  [S#] seguimiento personal\n"
        "  [B#] barbecho (labores en periodos sin proyecto)\n"
        "  [X#] evaluación extra fuera de proyecto (tema puntual, con justificación)\n\n"
        "REGLA DE TRAZABILIDAD (obligatoria, sin excepciones):\n"
        "TODA afirmación que escribas debe terminar con la etiqueta o etiquetas de las que proviene, "
        "p. ej. 'Entrega su trabajo a tiempo [E3][P2]'. "
        "PROHIBIDO escribir cualquier afirmación que no esté literalmente respaldada por una línea citada. "
        "Si no hay evidencia para una dimensión, escribe exactamente 'Sin información suficiente' (sin cita). "
        "No inventes, no infieras, no generalices más allá del texto citado. No uses etiquetas inexistentes.\n\n"
        "COBERTURA DE FUENTES (obligatoria):\n"
        "La regla anterior va de afirmación a fuente. Esta va al revés: TODA fuente debe acabar en uno "
        "de estos tres sitios, sin excepción:\n"
        "  (a) citada en CADA dimensión con cuyos CRITERIOS encaje su contenido —si una fuente habla de "
        "gestión de proyecto y de comunicación, cítala en las dos—, o\n"
        "  (b) recogida en 'sin_clasificar' si tiene contenido evaluable pero no encaja con los "
        "criterios de ninguna dimensión, o\n"
        "  (c) listada en '_fuentes_ignoradas' con su motivo, solo si no hay nada evaluable en ella.\n"
        "Antes de cerrar el JSON, repasa la lista de fuentes una a una y comprueba que ninguna se ha "
        "quedado fuera de (a), (b) y (c).\n"
        "Que una fuente sea genérica, poco concreta o repita lo que ya dice otra NO es motivo para "
        "descartarla: cítala en la dimensión cuyos criterios cubra.\n"
        "La cobertura NO te autoriza a meter una fuente en una dimensión con cuyos criterios no encaja: "
        "en ese caso va a 'sin_clasificar' —con su cita, como cualquier otro bullet—, nunca repartida a "
        "la fuerza. Es preferible una dimensión con 'Sin información suficiente' que una dimensión "
        "rellenada con evidencia que no le toca.\n\n"
        "EVOLUCIÓN TEMPORAL (importante):\n"
        "Los datos están en orden cronológico con su mes. NO promedies ni des una foto plana: describe la "
        "TRAYECTORIA a lo largo del año. No es lo mismo febrero que noviembre. Da MÁS PESO a lo más reciente, "
        "y cuando algo mejore o empeore, dilo citando ambos momentos (p. ej. 'empezó con poca autonomía en mar "
        "[E2] y cerró con mentalidad senior en nov [E9]'). Distingue también entre proyectos: un proyecto duro "
        "no es comparable a uno cómodo.\n\n"
        "BARBECHO: las labores de barbecho [B#] son, casi siempre, evidencia de 'contribution_to_firm' "
        "(contribución a la firma), no de las dimensiones de proyecto.\n\n"
        "EVALUACIONES EXTRA: las evaluaciones [X#] tratan temas puntuales fuera de proyecto; agrúpalas "
        "bajo 'evaluaciones_adicionales', no las fuerces en las dimensiones de proyecto salvo que el tema "
        "coincida claramente con una de ellas.\n\n"
        "Devuelve ÚNICAMENTE un JSON válido (sin bloques markdown) con esta estructura:\n"
        "{\n"
        '  "<clave_dimension>": {\n'
        '    "lider": "afirmación basada en evaluadores superiores [E1][P3]\\notra [E2]",\n'
        '    "equipo": "afirmación de iguales/subordinados [E4]",\n'
        '    "sin_nivel": "afirmación sin jerarquía clara [S1]"\n'
        "  },\n"
        "  ...\n"
        '  "contribution_to_firm": "bullets de contribución a la firma, cada uno con su cita [B1][O1]",\n'
        '  "evaluaciones_adicionales": "bullets de evaluaciones extra fuera de proyecto, cada uno con su cita [X1][X2]",\n'
        '  "sin_clasificar": "bullets con contenido evaluable que no encaja con los criterios de ninguna dimensión, cada uno con su cita [E7]",\n'
        '  "resultado": "valoración global en 2-3 frases que resuma la evolución del año",\n'
        '  "_fuentes_ignoradas": [{"cid": "E4", "motivo": "por qué no la has citado en ninguna dimensión"}]\n'
        "}\n\n"
        f"Dimensiones requeridas: {dims_lista}, contribution_to_firm, evaluaciones_adicionales, resultado.\n"
        "QUÉ DEFINE A CADA DIMENSIÓN (regla dura, por encima de la de cobertura):\n"
        "El mensaje del usuario trae, para cada dimensión, los criterios que la definen. Esos "
        "criterios —no su nombre, no su posición en la lista— deciden qué evidencia le "
        "corresponde. Aplica esta prueba a cada fuente y cada dimensión: '¿esta evidencia "
        "demuestra o contradice alguno de los criterios literales de esta dimensión?'. Si la "
        "respuesta es no, esa fuente NO va en esa dimensión, por muy relevante que sea la fuente "
        "y por mucho que el nombre de la dimensión sugiera lo contrario.\n"
        "Si los criterios de una dimensión están vacíos, son un texto de relleno o no describen "
        "ninguna conducta evaluable, entonces NINGUNA evidencia le corresponde: su valor es "
        "exactamente 'Sin información suficiente'. Que te queden fuentes sin colocar no es un "
        "error: van a 'sin_clasificar'. Rellenar una dimensión así SÍ es un error grave.\n"
        "Esta regla manda sobre la de cobertura y sobre cualquier impulso de dejar el informe "
        "completo. Las tres claves sin criterios (contribution_to_firm, evaluaciones_adicionales, "
        "resultado) quedan fuera de esta regla y se rigen por su significado habitual.\n"
        "Para cada dimensión agrupa: 'lider' = evaluadores superiores, "
        "'equipo' = mismo nivel o subordinados, 'sin_nivel' = sin jerarquía clara "
        "(coloca seguimiento personal y proyecto en el nivel que corresponda según quién evalúa). "
        "Dentro de una dimensión, omite los niveles que no tengan datos; la dimensión en sí "
        "siempre aparece, aunque sea con 'Sin información suficiente'. "
        "contribution_to_firm, evaluaciones_adicionales y resultado son cadenas planas, no objetos. "
        "resultado es una síntesis global de lo ya afirmado en las demás dimensiones: 2-3 frases "
        "en prosa, SIN citas [E#] y SIN meter evidencia nueva. No es una dimensión más.\n\n"
        "FORMATO DE LA RESPUESTA (obligatorio):\n"
        "Responde SOLO con el JSON. Tu primer carácter debe ser '{' y el último '}'. No "
        "escribas nada antes ni después: ni el repaso de las fuentes, ni tu razonamiento, "
        "ni notas sobre lo que has decidido. El repaso de cobertura hazlo mentalmente, no "
        "por escrito."
    )
    if idioma == "en":
        system += (
            "\n\nLANGUAGE: Write ALL comment text in English (the fields 'lider', 'equipo', "
            "'sin_nivel', 'contribution_to_firm', 'evaluaciones_adicionales' and 'resultado'). The source "
            "data may be in Spanish; translate the meaning, do not copy Spanish text. Keep the JSON keys "
            "and the citation tags like [E3] exactly as specified. When there is no evidence for a "
            "dimension, write exactly 'Not enough information' (without a citation)."
        )
    elif idioma == "pt":
        system += (
            "\n\nIDIOMA: Escreve TODO o texto dos comentários em português europeu (os campos "
            "'lider', 'equipo', 'sin_nivel', 'contribution_to_firm', 'evaluaciones_adicionales' "
            "e 'resultado'). Os dados de "
            "origem podem estar em espanhol; traduz o significado, não copies texto em espanhol. "
            "Mantém as chaves do JSON e as etiquetas de citação como [E3] exatamente como se indica. "
            "Quando não houver evidência para uma dimensão, escreve exatamente 'Informação "
            "insuficiente' (sem citação)."
        )

    system += config.INSTRUCCION_ANTIINYECCION

    user_content = (
        f"Empleado: {emp_data['empleado']}\n"
        f"Cargo: {cargo or 'No especificado'}\n"
        f"CA: {emp_data.get('ca', 'No especificado')}\n"
        f"{criterios_section}\n\n"
        f"{contexto}"
    )

    def _crear(system_arg):
        # El JSON lleva lider/equipo/sin_nivel por cada dimensión, más contribution,
        # evaluaciones adicionales, resultado y fuentes ignoradas: para un año entero, 4000
        # tokens se quedaban cortos. `max_tokens` es solo un techo (se paga lo generado),
        # así que se deja holgado. Con un techo alto hay que ir por streaming o la petición
        # agota el timeout HTTP del SDK; get_final_message() devuelve el mismo mensaje.
        with anthropic_client.messages.stream(
            model="claude-sonnet-4-6",
            max_tokens=16000,
            temperature=0,
            system=system_arg,
            messages=[{"role": "user", "content": user_content}],
        ) as stream:
            return stream.get_final_message()

    # El `system` (instrucciones + formato) es ESTÁTICO por (cargo, idioma) y se repite entre
    # empleados generados en ráfaga → se cachea (prompt caching): mismo modelo y misma calidad,
    # solo se paga una vez el prefijo durante la ventana de caché. La evidencia (variable por
    # persona) va en el mensaje del usuario, fuera de la caché.
    _t_ia = time.time()
    try:
        respuesta = _crear([{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}])
    except ErrorIA as err:
        if err.definitivo:
            # Sin saldo o API mal configurada: reintentar sin caché falla igual y gasta otra llamada.
            raise
        logging.warning("[informe anual] Prompt caching no disponible; reintento sin caché")
        respuesta = _crear(system)
    # [perf] para poder ver en el log si la espera está aquí o en las lecturas de Notion.
    logging.warning("[perf] redaccion del informe con Claude: %.1fs", time.time() - _t_ia)
    texto = "".join(b.text for b in respuesta.content if b.type == "text").strip()
    comentarios = _extraer_json_objeto(texto)
    if comentarios is None:
        # Antes esto era un json.loads pelado: el JSONDecodeError (que hereda de ValueError)
        # acababa en el handler de ValueError de la API y le pintaba al usuario el mensaje
        # crudo de Python. Se loguean principio y final: el fallo típico es texto de más.
        logging.error(
            "[informe anual] Respuesta no parseable de %s: stop_reason=%s, %d chars\n"
            "--- inicio ---\n%s\n--- final ---\n%s",
            emp_data.get("empleado", "?"), respuesta.stop_reason, len(texto),
            texto[:1500], texto[-1500:],
        )
        raise ErrorIA(
            "La IA no ha devuelto el análisis de las evaluaciones en el formato esperado. "
            "Vuelve a intentarlo; si sigue fallando, avisa al responsable de la "
            f"herramienta ({CONTACTO}).",
            codigo="informe_anual_respuesta_invalida",
        ) from None
    comentarios = _remapear_slugs_a_claves(comentarios, dims)
    comentarios = _validar_citas(comentarios, fuentes)
    _t_aud = time.time()
    comentarios["_avisos_verificacion"] = _verificar_soporte(comentarios, fuentes)
    logging.warning("[perf] auditoria de citas con Claude: %.1fs", time.time() - _t_aud)
    comentarios["_fuentes"] = fuentes
    return comentarios


# ── Word: helpers XML ─────────────────────────────────────────────────────────

def _dxb(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{lado}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tcPr.append(borders)


def _dxw(cell, inches):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def _dxr(para, texto, bold=False, size=9, underline=False, center=False):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(texto)
    run.bold = bold
    run.underline = underline
    run.font.name = "Outfit"
    run.font.size = Pt(size)
    return run


def _dxt(doc, texto):
    from docx.shared import Pt
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)
    _dxr(para, texto, bold=True, size=10, underline=True)
    return para


def _dx_hyperlink(para, url, texto, size=9):
    """Inserta un hyperlink real de Word (azul, subrayado) al final del párrafo."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r_id = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Outfit"); rFonts.set(qn("w:hAnsi"), "Outfit")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = texto
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)
    return hyperlink


def _dx_internal_link(para, anchor, texto, size=9):
    """Inserta un enlace interno de Word a un marcador (bookmark) del mismo documento."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Outfit"); rFonts.set(qn("w:hAnsi"), "Outfit")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = texto
    run.append(t)
    hyperlink.append(run)
    para._p.append(hyperlink)
    return hyperlink


def _dx_bookmark(para, name, bm_id):
    """Marca el párrafo con un bookmark para que los enlaces internos puedan saltar a él."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bm_id))
    para._p.insert(0, start)
    para._p.append(end)


def _dxr_con_citas(para, texto, fuentes=None, size=9):
    """Como _dxr pero convierte los tokens [E3]/[O1] en enlaces internos al anexo de Fuentes."""
    fuentes = fuentes or {}
    pos = 0
    for m in _CITE_RE.finditer(texto):
        if m.start() > pos:
            _dxr(para, texto[pos:m.start()], size=size)
        cid = m.group(1)
        if cid in fuentes:
            _dx_internal_link(para, f"fuente_{cid}", f"[{cid}]", size=size)
        else:
            _dxr(para, f"[{cid}]", size=size)
        pos = m.end()
    if pos < len(texto):
        _dxr(para, texto[pos:], size=size)


def _dx_bullets(cell, texto, fuentes=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    lineas = [l.strip(" •-–") for l in (texto or "").strip().splitlines() if l.strip()]
    if not lineas:
        return
    for i, linea in enumerate(lineas):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        if i == 0:
            p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _dxr(p, "• ", size=9)
        _dxr_con_citas(p, linea, fuentes, size=9)


def _dx_bullets_por_nivel(cell, contenido, fuentes=None, idioma="es"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if not isinstance(contenido, dict):
        _dx_bullets(cell, contenido, fuentes)
        return
    primer = True
    for nivel_key, label in _LABELS_NIVEL:
        texto = (contenido.get(nivel_key) or "").strip()
        if not texto:
            continue
        p = cell.paragraphs[0] if primer else cell.add_paragraph()
        if primer:
            p.clear()
        primer = False
        _dxr(p, _nivel_label(nivel_key, label, idioma) + ":", bold=True, size=8)
        for linea in texto.splitlines():
            linea = linea.strip(" •-–")
            if linea:
                pb = cell.add_paragraph()
                pb.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _dxr(pb, "• ", size=9)
                _dxr_con_citas(pb, linea, fuentes, size=9)


def _tabla_dims(doc, dims, comentarios, fuentes=None):
    tabla = doc.add_table(rows=len(dims) + 1, cols=3)
    tabla.style = "Table Grid"
    c0, c1, c2 = tabla.rows[0].cells
    for c, txt, w in ((c0, "Dimensión", _W_DIM), (c1, "Nota", _W_NOTA), (c2, "Comentarios del evaluador", _W_COM)):
        _dxb(c); _dxw(c, w)
        _dxr(c.paragraphs[0], txt, bold=True, size=9, center=(txt == "Nota"))
    for i, d in enumerate(dims):
        c0, c1, c2 = tabla.rows[i + 1].cells
        _dxb(c0); _dxw(c0, _W_DIM)
        _dxb(c1); _dxw(c1, _W_NOTA)
        _dxb(c2); _dxw(c2, _W_COM)
        _dxr(c0.paragraphs[0], d["etiqueta"], size=9)
        _dxr(c1.paragraphs[0], "X", size=9, center=True)
        _dx_bullets_por_nivel(c2, comentarios.get(d["clave"], ""), fuentes)
    return tabla


# ── HTML: generación ─────────────────────────────────────────────────────────

def guardar_informe_anual_html(emp_data: dict, comentarios: dict, cargo: str = "", idioma: str = "es",
                               incluir_fuentes: bool = True) -> str:
    # `incluir_fuentes=False` para todo documento que pueda acabar en manos del advisee: el anexo
    # trae el texto crudo de cada evaluación y el nombre del evaluador. Ver guardar_informe_anual_word.
    fuentes = comentarios.get("_fuentes", {}) if incluir_fuentes else {}

    def esc(v):
        return html_lib.escape(str(v or ""))

    def linkify(texto_html):
        """Convierte los tokens [E3]/[O1] (ya escapados) en anclas internas al anexo de Fuentes."""
        def repl(m):
            cid = m.group(1)
            src = fuentes.get(cid)
            if src:
                return (f"<a class='cite' href='#fuente-{cid}' "
                        f"title='{esc(src.get('label', ''))}'>[{cid}]</a>")
            return f"<span class='cite cite-off'>[{cid}]</span>"
        return _CITE_RE.sub(repl, texto_html)

    def bullets_html(texto):
        lineas = [ln.strip(" •-–") for ln in (texto or "").strip().splitlines() if ln.strip()]
        return "<br>".join(f"• {linkify(esc(ln))}" for ln in lineas) if lineas else "—"

    def bullets_html_por_nivel(contenido):
        if not isinstance(contenido, dict):
            return bullets_html(contenido)
        partes = []
        for nivel_key, label in _LABELS_NIVEL:
            texto = (contenido.get(nivel_key) or "").strip()
            if not texto:
                continue
            lineas = [ln.strip(" •-–") for ln in texto.splitlines() if ln.strip()]
            buls = "<br>".join(f"• {linkify(esc(ln))}" for ln in lineas)
            partes.append(f"<span style='font-size:11px;font-weight:700'>{esc(_nivel_label(nivel_key, label, idioma))}:</span><br>{buls}")
        return "<br><br>".join(partes) if partes else "—"

    def filas_dims(dims):
        filas = ""
        for d in dims:
            etiqueta = esc(_dim_label(d["slug"], d["etiqueta"], idioma))
            filas += f"<tr><td>{etiqueta}</td><td class='nc'>X</td><td>{bullets_html_por_nivel(comentarios.get(d['clave'],''))}</td></tr>"
        return filas

    dims_proyectos = dimensiones_informe(emp_data.get("empleado", ""), cargo, incluir_liderazgo=False)

    cargo_lower = cargo.strip().lower()
    requiere_liderazgo = any(c in cargo_lower for c in _REQUIERE_LIDERAZGO)

    cargo_row = f"<tr><td><strong>{t('anual.role', idioma)}</strong></td><td>{esc(cargo)}</td></tr>" if cargo else ""

    liderazgo_bloque = ""
    if requiere_liderazgo:
        liderazgo_bloque = f"""
        <h2 class="sec">{t("anual.leadership", idioma)}</h2>
        <table class="et"><thead><tr><th>{t("anual.col_dimension", idioma)}</th><th class="nc">{t("anual.col_score", idioma)}</th><th>{t("anual.col_eval_comments", idioma)}</th></tr></thead>
        <tbody>{filas_dims(_dims_fijas(_DIMS_LIDERAZGO))}</tbody></table>"""

    evals_adicionales_bloque = ""
    if emp_data.get("evals_extra"):
        evals_adicionales_bloque = f"""
        <h2 class="sec">{t("anual.additional_evals", idioma)}</h2>
        <p>{bullets_html(comentarios.get('evaluaciones_adicionales',''))}</p>"""

    # Evidencia con contenido pero que no encaja con los criterios de ninguna dimensión.
    # Se muestra aparte —"cosas que la IA no ha sabido clasificar"— en vez de forzarla en
    # una dimensión: así no se pierde y el CA ve qué se ha quedado fuera.
    sin_clasificar_bloque = ""
    if (comentarios.get("sin_clasificar") or "").strip():
        sin_clasificar_bloque = f"""
        <h2 class="sec">{t("anual.unclassified", idioma)}</h2>
        <p class="fine">{t("anual.unclassified_hint", idioma)}</p>
        <p>{bullets_html(comentarios.get('sin_clasificar',''))}</p>"""

    objetivos_html = ""
    objetivos = emp_data.get("objetivos", [])
    if objetivos:
        items_html = ""
        for obj in objetivos:
            titulo_o = esc(obj.get("titulo", ""))
            tipo_o = esc(obj.get("tipo", ""))
            kpis_o = esc(obj.get("kpis", ""))
            desc_o = esc(obj.get("descripcion", ""))
            ca_o = esc(obj.get("ca", ""))
            fecha_o = esc((obj.get("fecha") or "")[:10])
            header = f"<strong>{titulo_o}</strong>"
            if tipo_o:
                header += f" <span class='fine'>({tipo_o})</span>"
            meta = f"<span class='fine'>{ca_o} — {fecha_o}</span>" if (ca_o or fecha_o) else ""
            kpis_block = f"<p><em>KPIs:</em> {kpis_o}</p>" if kpis_o else ""
            desc_block = f"<p>{desc_o}</p>" if desc_o else ""
            items_html += f"<div style='margin-bottom:12px'><p>{header}</p>{meta}{kpis_block}{desc_block}</div>"
        objetivos_html = items_html
    else:
        objetivos_html = f"<p>{t('anual.no_goals', idioma)}</p>"

    # Panel de revisión (solo borrador): avisos del verificador + bullets descartados.
    # El CA decide qué hacer antes de publicar el informe final.
    avisos = comentarios.get("_avisos_verificacion", [])
    descartados = comentarios.get("_bullets_descartados", [])
    panel_revision = ""
    if avisos or descartados:
        partes_panel = []
        if avisos:
            items = ""
            for a in avisos:
                motivo = esc(a.get("motivo", ""))
                items += (f"<li>{linkify(esc(a.get('afirmacion', '')))}"
                          f"{f' — <em>{motivo}</em>' if motivo else ''}</li>")
            partes_panel.append(
                f"<p class='rev-h'>{t('anual.rev_warn_title', idioma)}</p>"
                f"<ul>{items}</ul>"
            )
        if descartados:
            items = "".join(f"<li>{esc(d)}</li>" for d in descartados)
            partes_panel.append(
                f"<p class='rev-h'>{t('anual.rev_discarded_title', idioma)}</p>"
                f"<ul>{items}</ul>"
            )
        panel_revision = (
            "<div class='revision'>"
            f"<p class='rev-title'>{t('anual.rev_title', idioma)}</p>"
            f"<p class='rev-sub'>{t('anual.rev_sub', idioma)}</p>"
            + "".join(partes_panel) +
            "</div>"
        )

    # Anexo "Fuentes / Evidencia": cada cita del informe enlaza aquí (sin depender de Notion).
    _TIPO_LABEL = {
        "opinion": t("anual.src_opinion", idioma), "evaluacion": t("anual.src_evaluacion", idioma),
        "proyecto": t("anual.src_proyecto", idioma), "seguimiento": t("anual.src_seguimiento", idioma),
        "barbecho": t("anual.src_barbecho", idioma), "extra": t("anual.src_extra", idioma),
        "aportacion_ca": t("anual.src_aportacion_ca", idioma),
    }
    _ORDEN_TIPO = {"O": 0, "E": 1, "P": 2, "S": 3, "B": 4, "X": 5, "C": 6}

    def _sort_fuente(cid):
        return (_ORDEN_TIPO.get(cid[:1], 9), int(cid[1:]) if cid[1:].isdigit() else 0)

    fuentes_items = ""
    for cid in sorted(fuentes.keys(), key=_sort_fuente):
        src = fuentes[cid]
        evaluador = src.get("evaluador", "")
        meta = f" · <span class='f-eval'>{esc(evaluador)}</span>" if evaluador else ""
        fuentes_items += (
            f"<div class='fuente' id='fuente-{cid}'>"
            f"<span class='f-id'>[{cid}]</span> "
            f"<span class='f-tipo'>{esc(_TIPO_LABEL.get(src.get('tipo'), ''))}</span> · "
            f"<span class='f-label'>{esc(src.get('label', ''))}</span>{meta}"
            f"<p class='f-texto'>{esc(src.get('texto', '')) or '—'}</p>"
            f"</div>"
        )
    fuentes_html = (
        f"<h2 class='sec'>{t('anual.sources_evidence', idioma)}</h2>"
        f"<p class='f-intro'>{t('anual.sources_intro_web', idioma)}</p>"
        f"{fuentes_items}"
    ) if fuentes else ""

    fecha = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    # Mismo criterio que el Word: se evalúa el año anterior al de generación.
    año = datetime.now(timezone.utc).year - 1

    contenido = f"""<!DOCTYPE html>
<html lang="{idioma}">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{t("anual.title_web", idioma, emp=esc(emp_data['empleado']))}</title>
<style>
{config.IGENERIS_CSS}
.shell {{ max-width: 960px; margin: 0 auto; padding-bottom: 60px; }}
.top {{ padding-top: clamp(42px, 8vw, 92px); margin-bottom: 36px; }}
.it {{ width: 100%; border-collapse: collapse; margin-bottom: 24px; font-size: 14px; }}
.it td {{ border: 1px solid var(--ink); padding: 8px 14px; }}
.et {{ width: 100%; border-collapse: collapse; margin-bottom: 24px; font-size: 14px; }}
.et th, .et td {{ border: 1px solid var(--ink); padding: 8px 12px; vertical-align: top; text-align: justify; }}
.et th {{ background: var(--soft); font-weight: 700; font-size: 12px; text-transform: uppercase; letter-spacing: .05em; }}
.et td:first-child {{ width: 200px; font-weight: 500; }}
.nc {{ width: 60px; text-align: center; }}
.sec {{ font-size: 13px; text-transform: uppercase; letter-spacing: .08em; border-bottom: 2px solid var(--ink); padding-bottom: 6px; margin: 32px 0 14px; color: var(--ink); }}
.rg {{ display: grid; grid-template-columns: 130px 1fr; border: 1px solid var(--ink); font-size: 14px; }}
.rg > div {{ padding: 14px 16px; }}
.rg > div:first-child {{ border-right: 1px solid var(--ink); text-align: center; font-weight: 700; }}
.cite {{ font-size: 10px; font-weight: 600; color: #0563C1; text-decoration: none; vertical-align: super; padding: 0 1px; }}
.cite:hover {{ text-decoration: underline; }}
.cite-off {{ color: #999; cursor: help; }}
.revision {{ border: 1px solid #d9a300; background: #fff8e6; border-radius: 8px; padding: 16px 20px; margin-bottom: 28px; font-size: 13px; }}
.revision .rev-title {{ font-weight: 700; font-size: 14px; margin: 0 0 2px; color: #8a6d00; }}
.revision .rev-sub {{ margin: 0 0 12px; color: #8a6d00; font-size: 12px; }}
.revision .rev-h {{ font-weight: 700; margin: 12px 0 4px; }}
.revision ul {{ margin: 0 0 4px; padding-left: 20px; }}
.revision li {{ margin-bottom: 4px; }}
.f-intro {{ font-size: 13px; color: #555; margin-bottom: 16px; }}
.fuente {{ border-left: 3px solid var(--ink); padding: 8px 14px; margin-bottom: 12px; font-size: 13px; scroll-margin-top: 80px; }}
.fuente:target {{ background: #fff3cd; border-left-color: #d9a300; }}
.f-id {{ font-weight: 700; color: #0563C1; }}
.f-tipo {{ font-weight: 600; }}
.f-eval {{ font-style: italic; }}
.f-texto {{ margin: 6px 0 0; white-space: pre-line; }}
</style>
</head>
<body>
<main class="page shell">
<nav class="nav">
  <a class="brand" href="javascript:void(0)" onclick="window.close()">igeneris</a>
  <div class="nav-links"><button class="secondary" onclick="window.close()">{t("report.cerrar", idioma)}</button></div>
</nav>
<div class="top">
  <p class="kicker">{t("anual.eval_year", idioma, anio=año)}</p>
  <h1>{esc(emp_data['empleado'])}</h1>
  <p>{t("anual.generated", idioma, fecha=fecha)}</p>
</div>

{panel_revision}

<table class="it">
  <tr><td><strong>{t("anual.name", idioma)}</strong></td><td>{esc(emp_data['empleado'])}</td></tr>
  {cargo_row}
  <tr><td><strong>Career Advisor</strong></td><td>{esc(emp_data.get('ca') or '—')}</td></tr>
</table>

<h2 class="sec">{t("anual.rating_year", idioma, anio=f"{año}/{año + 1}")}</h2>
<table class="et">
  <thead><tr><th>{t("anual.col_dimension", idioma)}</th><th class="nc">{t("anual.col_score", idioma)}</th><th>{t("anual.col_eval_comments", idioma)}</th></tr></thead>
  <tbody>{filas_dims(dims_proyectos)}</tbody>
</table>

{liderazgo_bloque}

<h2 class="sec">CONTRIBUTION TO THE FIRM</h2>
<p>{bullets_html(comentarios.get('contribution_to_firm',''))}</p>

{evals_adicionales_bloque}
{sin_clasificar_bloque}

<h2 class="sec">{t("anual.result", idioma)}</h2>
<div class="rg">
  <div>{t("anual.overall_score", idioma)}<br><strong>X / 5</strong></div>
  <div>{esc(comentarios.get('resultado','—'))}</div>
</div>

<h2 class="sec">{t("anual.objectives_year", idioma, anio=año + 1)}</h2>
{objetivos_html}

{fuentes_html}
</main>
</body>
</html>"""

    os.makedirs(config.CARPETA_WEB, exist_ok=True)
    slug = slug_archivo(emp_data["empleado"])
    ruta = os.path.join(config.CARPETA_WEB, f"informe_anual_{slug}.html")
    with open(ruta, "w", encoding="utf-8") as f:
        f.write(contenido)
    logging.info("Informe anual HTML guardado: %s", ruta)
    return slug


# ── Word: generación ─────────────────────────────────────────────────────────

def _celda_emp(cell, label, valor, w):
    """Rellena una celda de la cabecera: etiqueta bold o valor normal."""
    _dxb(cell); _dxw(cell, w)
    if label:
        _dxr(cell.paragraphs[0], label, bold=True, size=9)
    else:
        _dxr(cell.paragraphs[0], valor or "", size=9)


def guardar_informe_anual_word(emp_data: dict, comentarios: dict, cargo: str = "", idioma: str = "es",
                               valores_ca: dict | None = None, nombre_archivo: str = "",
                               incluir_fuentes: bool = True) -> str:
    """Genera el .docx replicando la plantilla oficial de EVALUACIÓN ANUAL de IGENERIS.

    Campos que el sistema no posee (CA '26, salarios, % variable, promoción, deadlines,
    y la NOTA por dimensión) se dejan en blanco para que el CA los rellene, igual que en
    la plantilla. Los comentarios por dimensión sí los redacta Claude, con sus citas a Notion.

    `valores_ca`: valores de esos huecos rellenados por el CA en el borrador web
    (claves: caSiguiente, salarioActual, notas{clave}, retribucion{...}, resultadoEval{...},
    objetivos[{texto,deadline}]). Si falta una clave, su hueco se queda en blanco.
    `nombre_archivo`: nombre del .docx a escribir en CARPETA_WEB (por defecto
    informe_anual_{slug}.docx).
    `incluir_fuentes`: el anexo de Fuentes/Evidencia es SOLO para el CA. Lleva el texto literal
    de cada evaluación y el nombre de quien la escribió, así que enseñárselo al advisee rompe el
    anonimato de los evaluadores. Pásalo en False en cualquier documento que él pueda descargar.
    """
    if Document is None:
        raise RuntimeError("Instala python-docx: pip install python-docx")

    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    vca = valores_ca or {}
    notas_ca = vca.get("notas") or {}
    ret_ca = vca.get("retribucion") or {}
    res_ca = vca.get("resultadoEval") or {}

    def _v(d, clave):
        return str(d.get(clave) or "").strip()

    fuentes = comentarios.get("_fuentes", {}) if incluir_fuentes else {}
    ahora = datetime.now(timezone.utc)
    anio_eval = ahora.year - 1          # se evalúa el año anterior (p. ej. 2025 en marzo 2026)
    anio_sig = ahora.year               # año siguiente al evaluado
    fecha_txt = f"{_mes_label(ahora.month - 1, idioma)} {ahora.year}"
    yy = f"'{str(anio_eval)[-2:]}"      # "'25"
    yy_sig = f"'{str(anio_sig)[-2:]}"   # "'26"

    doc = Document()
    sec = doc.sections[0]
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.76))

    # ── Marca + título ────────────────────────────────────────────────────────
    marca = doc.add_paragraph()
    marca.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dxr(marca, ".Igeneris", bold=True, size=22)
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _dxr(titulo, t("anual.doc_title", idioma), bold=True, size=13, underline=True)
    doc.add_paragraph()

    # ── Tabla de datos del empleado (4 columnas) ──────────────────────────────
    w1, w2, w3, w4 = 1.2, 2.24, 1.3, _CONTENT_W_IN - 1.2 - 2.24 - 1.3
    filas_emp = [
        (t("anual.employee", idioma), emp_data["empleado"], t("anual.date", idioma), fecha_txt),
        (f"CA {yy}", emp_data.get("ca", ""), t("anual.current_position", idioma), cargo or ""),
        (f"CA {yy_sig}", _v(vca, "caSiguiente"), t("anual.current_salary", idioma), _v(vca, "salarioActual")),
        # El área decide qué criterios y qué apartados lleva el informe, así que conviene
        # que quede escrita en el propio documento.
        (t("anual.area", idioma), _grupo_empleado(emp_data.get("empleado", ""), cargo), "", ""),
    ]
    t_emp = doc.add_table(rows=len(filas_emp), cols=4)
    t_emp.style = "Table Grid"
    for i, (l1, v1, l2, v2) in enumerate(filas_emp):
        c0, c1, c2, c3 = t_emp.rows[i].cells
        _celda_emp(c0, l1, None, w1)
        _celda_emp(c1, None, v1, w2)
        _celda_emp(c2, l2, None, w3)
        _celda_emp(c3, None, v2, w4)
    doc.add_paragraph()

    # ── CALIFICACIÓN {año_eval}/{año_sig} ─────────────────────────────────────
    _dxt(doc, t("anual.rating_year", idioma, anio=f"{anio_eval}/{anio_sig}"))

    dims = dimensiones_informe(emp_data.get("empleado", ""), cargo)

    w_dim, w_nota = 1.6, 0.6
    w_com = _CONTENT_W_IN - w_dim - w_nota
    t_cal = doc.add_table(rows=len(dims) + 1, cols=3)
    t_cal.style = "Table Grid"
    h0, h1, h2 = t_cal.rows[0].cells
    for c, txt, w, ctr in ((h0, t("anual.projects", idioma), w_dim, False), (h1, t("anual.score_up", idioma), w_nota, True), (h2, t("anual.comments_up", idioma), w_com, False)):
        _dxb(c); _dxw(c, w)
        _dxr(c.paragraphs[0], txt, bold=True, size=9, center=ctr)
    for i, d in enumerate(dims):
        clave = d["clave"]
        c0, c1, c2 = t_cal.rows[i + 1].cells
        _dxb(c0); _dxw(c0, w_dim)
        _dxb(c1); _dxw(c1, w_nota)
        _dxb(c2); _dxw(c2, w_com)
        _dxr(c0.paragraphs[0], _dim_label(d["slug"], d["etiqueta"], idioma), size=9)
        # NOTA: en blanco salvo que el CA la haya rellenado en el borrador web
        nota_dim = _v(notas_ca, clave)
        if nota_dim:
            _dxr(c1.paragraphs[0], nota_dim, size=9, center=True)
        _dx_bullets_por_nivel(c2, comentarios.get(clave, ""), fuentes, idioma=idioma)
    doc.add_paragraph()

    # ── Notas finales / retribución ───────────────────────────────────────────
    wc1, wc2, wc3, wc4 = 2.9, 0.7, 1.5, _CONTENT_W_IN - 2.9 - 0.7 - 1.5
    total_var = t("anual.total_variable", idioma, yy=yy)
    if _v(ret_ca, "totalVariable"):
        total_var += f" {_v(ret_ca, 'totalVariable')}"
    filas_ret = [
        (t("anual.final_projects", idioma), _v(ret_ca, "notaProyectos"), t("anual.variable_60", idioma), _v(ret_ca, "variable60")),
        (t("anual.final_contrib", idioma), _v(ret_ca, "notaContribucion"), t("anual.variable", idioma), _v(ret_ca, "variable")),
        (t("anual.corp_objectives", idioma), _v(ret_ca, "objetivosCorporativos"), t("anual.variable_30", idioma), total_var),
    ]
    t_ret = doc.add_table(rows=len(filas_ret), cols=4)
    t_ret.style = "Table Grid"
    for i, (l1, v1, l2, v2) in enumerate(filas_ret):
        c0, c1, c2, c3 = t_ret.rows[i].cells
        for c, w in ((c0, wc1), (c1, wc2), (c2, wc3), (c3, wc4)):
            _dxb(c); _dxw(c, w)
        _dxr(c0.paragraphs[0], l1, bold=True, size=9)
        _dxr(c1.paragraphs[0], v1, size=9, center=True)
        _dxr(c2.paragraphs[0], l2, size=9)
        _dxr(c3.paragraphs[0], v2, bold=(i == len(filas_ret) - 1 and bool(v2)), size=9)
    doc.add_paragraph()

    # ── RESULTADO EVAL {año} ──────────────────────────────────────────────────
    _dxt(doc, t("anual.eval_result", idioma, yy=yy))
    wr = [1.2, 0.9, 1.4, 1.3, _CONTENT_W_IN - 1.2 - 0.9 - 1.4 - 1.3]
    t_res = doc.add_table(rows=1, cols=5)
    t_res.style = "Table Grid"
    cells = t_res.rows[0].cells
    salario_nuevo = t("anual.new_fixed_salary", idioma)
    if _v(res_ca, "nuevoSalarioFijo"):
        salario_nuevo += f" {_v(res_ca, 'nuevoSalarioFijo')}"
    textos_res = [(t("anual.promotion", idioma), True, False), (_v(res_ca, "promocion"), False, True),
                  (t("anual.position_next", idioma, yy=yy_sig), True, False),
                  (_v(res_ca, "cargoSiguiente"), False, True), (salario_nuevo, True, False)]
    for c, w, (txt, bold, ctr) in zip(cells, wr, textos_res):
        _dxb(c); _dxw(c, w)
        _dxr(c.paragraphs[0], txt, bold=bold, size=9, center=ctr)
    doc.add_paragraph()

    # ── OPORTUNIDADES DE MEJORA / OBJETIVOS {año+1} ───────────────────────────
    _dxt(doc, t("anual.improvement_objectives", idioma, yy=yy_sig))
    objetivos = emp_data.get("objetivos", [])
    objetivos_ca = vca.get("objetivos")  # [{texto, deadline}] editados en el borrador web
    n_filas = max(3, len(objetivos_ca) if objetivos_ca is not None else len(objetivos))
    w_obj, w_dl = _CONTENT_W_IN - 1.2, 1.2
    t_obj = doc.add_table(rows=n_filas + 1, cols=2)
    t_obj.style = "Table Grid"
    ch0, ch1 = t_obj.rows[0].cells
    _dxb(ch0); _dxw(ch0, w_obj)
    _dxb(ch1); _dxw(ch1, w_dl)
    _dxr(ch0.paragraphs[0], "", size=9)
    _dxr(ch1.paragraphs[0], t("anual.deadline", idioma), bold=True, size=9, center=True)
    for i in range(n_filas):
        c0, c1 = t_obj.rows[i + 1].cells
        _dxb(c0); _dxw(c0, w_obj)
        _dxb(c1); _dxw(c1, w_dl)
        texto_obj, deadline_obj = "", ""
        if objetivos_ca is not None:
            if i < len(objetivos_ca):
                texto_obj = str(objetivos_ca[i].get("texto") or "").strip()
                deadline_obj = _fmt_deadline(objetivos_ca[i].get("deadline"))
        elif i < len(objetivos):
            o = objetivos[i]
            texto_obj = o.get("titulo") or o.get("descripcion") or ""
        _dxr(c0.paragraphs[0], f"{i + 1}.  {texto_obj}".rstrip(), size=9)
        if deadline_obj:
            _dxr(c1.paragraphs[0], deadline_obj, size=9, center=True)

    # ── FUENTES / EVIDENCIA (anexo, cada cita salta aquí) ─────────────────────
    if fuentes:
        _ORDEN = {"O": 0, "E": 1, "P": 2, "S": 3, "B": 4, "X": 5, "C": 6}
        _TIPO = {
            "opinion": t("anual.src_opinion", idioma), "evaluacion": t("anual.src_evaluacion", idioma),
            "proyecto": t("anual.src_proyecto", idioma), "seguimiento": t("anual.src_seguimiento", idioma),
            "barbecho": t("anual.src_barbecho", idioma), "extra": t("anual.src_extra", idioma),
            "aportacion_ca": t("anual.src_aportacion_ca", idioma),
        }
        doc.add_paragraph()
        _dxt(doc, t("anual.sources_evidence", idioma))
        _dxr(doc.add_paragraph(),
             t("anual.sources_intro_docx", idioma),
             size=8)
        ordenadas = sorted(fuentes, key=lambda c: (_ORDEN.get(c[:1], 9), int(c[1:]) if c[1:].isdigit() else 0))
        for n, cid in enumerate(ordenadas, 1):
            src = fuentes[cid]
            p = doc.add_paragraph()
            _dx_bookmark(p, f"fuente_{cid}", n)
            _dxr(p, f"[{cid}] ", bold=True, size=9)
            _dxr(p, f"{_TIPO.get(src.get('tipo'), '')} · {src.get('label', '')}", bold=True, size=9)
            if src.get("evaluador"):
                _dxr(p, f" · {src['evaluador']}", size=9)
            _dxr(doc.add_paragraph(), src.get("texto", "") or "—", size=9)

    # ── Guardar ───────────────────────────────────────────────────────────────
    os.makedirs(config.CARPETA_WEB, exist_ok=True)
    slug = slug_archivo(emp_data["empleado"])
    ruta = os.path.join(config.CARPETA_WEB, nombre_archivo or f"informe_anual_{slug}.docx")
    doc.save(ruta)
    logging.info("Informe anual guardado: %s", ruta)
    return slug


# ── Caché ─────────────────────────────────────────────────────────────────────

def _huella_datos(emp_data: dict, cargo: str = "", criterios: str = "") -> str:
    datos = {
        "v": 5,
        "opiniones": emp_data.get("opiniones_ca", []),
        "evaluaciones": emp_data.get("evaluaciones", []),
        "evals_proyecto": emp_data.get("evals_proyecto", []),
        "seguimiento": emp_data.get("seguimiento", []),
        "barbecho": emp_data.get("barbecho", []),
        "evals_extra": emp_data.get("evals_extra", []),
        "cargo": (cargo or "").strip().lower(),
        # Criterios DTI de Notion (varían por grupo: Negocio/MiddleOffice/Palantir).
        # Si cambian en Notion, la huella cambia y el informe se regenera.
        "criterios": criterios or "",
    }
    return hashlib.sha256(
        json.dumps(datos, ensure_ascii=False, sort_keys=True).encode()
    ).hexdigest()


def _ruta_cache(slug: str) -> str:
    return os.path.join(config.CARPETA_WEB, f"informe_anual_{slug}_cache.json")


def _leer_cache(slug: str) -> dict | None:
    ruta = _ruta_cache(slug)
    if not os.path.exists(ruta):
        return None
    try:
        with open(ruta, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _escribir_cache(slug: str, huella: str) -> None:
    os.makedirs(config.CARPETA_WEB, exist_ok=True)
    with open(_ruta_cache(slug), "w", encoding="utf-8") as f:
        json.dump({"huella": huella}, f)


# ── Punto de entrada ──────────────────────────────────────────────────────────

def generar_informe_anual(evaluado: str, cargo: str = "") -> str:
    """Lee Notion → interpreta con Claude → genera Word. Reutiliza caché si no hay cambios."""
    emp_data = obtener_datos_empleado_anual(evaluado)
    if not emp_data.get("opiniones_ca") and not emp_data.get("evaluaciones"):
        raise ValueError(
            f"No hay opiniones del CA ni evaluaciones mensuales para '{evaluado}'."
        )

    slug = slug_archivo(evaluado)
    idioma = idioma_de_persona(evaluado)
    # Se computa una sola vez: alimenta la huella de caché y el prompt (evita doble lectura de Notion).
    # Las dimensiones entran en la huella: si en Notion se añade, quita, renombra o
    # reordena un criterio, la huella cambia y el informe se regenera con la plantilla
    # nueva. Mientras no cambien, se reutiliza el de siempre y no se gasta IA.
    dims = dimensiones_informe(evaluado, cargo)
    criterios = _criterios_para_prompt(cargo, idioma, evaluado, dims=dims)
    huella = _huella_datos(
        emp_data, cargo=cargo,
        criterios=(criterios or "") + f"|lang={idioma}|dims={huella_dimensiones(dims)}",
    )
    ruta_docx = os.path.join(config.CARPETA_WEB, f"informe_anual_{slug}.docx")
    ruta_html = os.path.join(config.CARPETA_WEB, f"informe_anual_{slug}.html")
    cache = _leer_cache(slug)

    if cache and cache.get("huella") == huella and os.path.exists(ruta_docx) and os.path.exists(ruta_html):
        logging.info("Informe anual en caché para %s, reutilizando.", evaluado)
        return slug

    comentarios = interpretar_evaluaciones_anual(emp_data, cargo=cargo, criterios=criterios,
                                                 idioma=idioma, dims=dims)
    slug = guardar_informe_anual_word(emp_data, comentarios, cargo=cargo, idioma=idioma)
    guardar_informe_anual_html(emp_data, comentarios, cargo=cargo, idioma=idioma)
    _escribir_cache(slug, huella)
    return slug
